#!/usr/bin/env python3
"""Create IMGC PAS Application Guide using python-docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

SCREENSHOTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')
OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'IMGC_PAS_Application_Guide.docx')

# Brand Colors
NAVY_HEX = '1E3A5F'
ORANGE_HEX = 'F59E0B'
DARK_TEXT_HEX = '1E293B'
GRAY_HEX = '64748B'
LIGHT_BG_HEX = 'FFF7ED'
WHITE_HEX = 'FFFFFF'
TEAL_HEX = '0D9488'
GREEN_HEX = '10B981'
RED_HEX = 'EF4444'

doc = Document()

# ============================================================================
# STYLES
# ============================================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    heading_style.font.bold = True
    if level == 1:
        heading_style.font.size = Pt(24)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(18)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading_style.font.size = Pt(14)
        heading_style.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_orange_heading(text, level=2):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
    return p


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{NAVY_HEX}"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = LIGHT_BG_HEX if r_idx % 2 == 0 else WHITE_HEX
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            p = row_cells[c_idx].paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def add_screenshot(filename, width=6.0):
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(filepath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=Inches(width))
        return True
    else:
        doc.add_paragraph(f'[Screenshot: {filename} - not found]')
        return False


def add_bullet_list(items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'


def add_note_box(text, prefix="Note"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{prefix}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
    run = p.add_run(text)
    run.font.size = Pt(10)


# ============================================================================
# COVER PAGE
# ============================================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('IMGC CORPORATION')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Policy Administration System')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Application Guide')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
run.font.name = 'Calibri'

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Enterprise Edition  |  Version 3.2.1')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('March 2026  |  Confidential')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.font.name = 'Calibri'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('India Mortgage Guarantee Corporation')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('End-to-End Insurance Operations Management Platform')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('1.', 'Platform Overview', '3'),
    ('2.', 'Login & Authentication', '4'),
    ('3.', 'Executive Dashboard', '5'),
    ('4.', 'Policy Management', '7'),
    ('5.', 'Policy Detail & Lifecycle', '8'),
    ('6.', 'Customer Management', '10'),
    ('7.', 'Underwriting Engine', '11'),
    ('8.', 'Claims Management', '13'),
    ('9.', 'Billing & Payments', '15'),
    ('10.', 'Task Management', '17'),
    ('11.', 'Document Management', '18'),
    ('12.', 'Reports & Analytics', '19'),
    ('13.', 'Administration & User Management', '20'),
    ('14.', 'Compliance Management', '21'),
    ('15.', 'Technology Stack & Architecture', '22'),
    ('16.', 'User Roles & Permissions', '23'),
    ('17.', 'Navigation Reference', '24'),
    ('18.', 'API Endpoints Reference', '25'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  ')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
    run.font.bold = True
    run = p.add_run(f'{title}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_page_break()

# ============================================================================
# 1. PLATFORM OVERVIEW
# ============================================================================
doc.add_heading('1. Platform Overview', level=1)

doc.add_paragraph(
    'The IMGC Policy Administration System (PAS) is an enterprise-grade insurance operations '
    'management platform developed for India Mortgage Guarantee Corporation. It provides '
    'end-to-end policy lifecycle management, claims processing, underwriting, billing, '
    'compliance tracking, and comprehensive reporting capabilities.'
)

doc.add_heading('Platform Highlights', level=2)

add_styled_table(
    ['Feature', 'Details'],
    [
        ['Modules', '14+ Enterprise Modules'],
        ['API Endpoints', '150+ RESTful APIs with JWT Authentication'],
        ['User Roles', '5 Role-Based Access Levels (Admin, Underwriter, Claims, Agent, Auditor)'],
        ['Security', '256-bit Encryption, SOC 2 Compliant, IRDAI Regulated'],
        ['Frontend', 'Next.js 14, React 18, TypeScript, Tailwind CSS'],
        ['Backend', 'Express.js, TypeScript, JWT Auth with Refresh Tokens'],
        ['Charts', 'Recharts — Interactive bar, pie, donut, horizontal bar charts'],
        ['Design', 'Enterprise orange accent theme (30% coverage), responsive design'],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_heading('Core Capabilities', level=2)

add_bullet_list([
    'Policy Lifecycle Management — Quote to issuance, endorsement, renewal, and cancellation',
    'Claims Automation — FNOL intake through settlement tracking with reserve management',
    'Risk Intelligence — Real-time underwriting scoring with configurable rules engine',
    'Billing & Payments — Full invoicing, installment plans, receivables aging, and ledger',
    'Task Management — SLA-tracked task queues with workflow-driven auto-assignment',
    'Document Management — Upload, template-based generation, categorization, versioning',
    'Compliance & Audit — Regulatory requirement tracking with complete audit trails',
    'Executive Analytics — 7 KPIs, 5+ chart types, role-based dashboards, date-filtered reports',
])

doc.add_page_break()

# ============================================================================
# 2. LOGIN & AUTHENTICATION
# ============================================================================
doc.add_heading('2. Login & Authentication', level=1)

add_screenshot('01-login.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The login page features a split-screen enterprise design. The left panel showcases '
    'the IMGC brand with platform capabilities (Policy Lifecycle, Claims Automation, '
    'Risk Intelligence, Compliance Ready). The right panel provides a clean login form '
    'with email and password fields.'
)

doc.add_heading('Authentication Features', level=3)
add_bullet_list([
    'JWT-based authentication with access and refresh token mechanism',
    'Encrypted password storage using bcrypt hashing algorithm',
    'Automatic session management with configurable token expiry',
    'Role-based access control (RBAC) — permissions enforced from login onward',
    '256-bit end-to-end encryption for all data in transit',
    'SOC 2 compliance badge and IRDAI regulatory compliance indicators on login',
])

doc.add_heading('Security Badges', level=3)
doc.add_paragraph(
    'The login page displays three compliance indicators: 256-bit encrypted, '
    'SOC 2 compliant, and IRDAI regulated. Platform statistics are shown — '
    '14+ Modules, 150+ API Endpoints, and 5 User Roles.'
)

doc.add_heading('Login Credentials (Demo)', level=3)
add_styled_table(
    ['Role', 'Email', 'Password'],
    [
        ['Admin', 'admin@pas.com', 'password123'],
        ['Underwriter', 'underwriter@pas.com', 'password123'],
        ['Claims Adjuster', 'claims@pas.com', 'password123'],
        ['Agent', 'agent@pas.com', 'password123'],
        ['Auditor', 'auditor@pas.com', 'password123'],
    ],
    col_widths=[1.5, 2.5, 1.5]
)

doc.add_page_break()

# ============================================================================
# 3. EXECUTIVE DASHBOARD
# ============================================================================
doc.add_heading('3. Executive Dashboard', level=1)

add_screenshot('02-dashboard-top.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Executive Dashboard provides a comprehensive real-time overview of the entire '
    'insurance operation. It features role-based views that adapt based on the logged-in '
    'user\'s role. The default Executive View displays 7 KPIs with target tracking, '
    '4 summary cards, and 5+ analytics charts.'
)

doc.add_heading('Executive KPIs', level=3)
doc.add_paragraph(
    'Seven key performance indicators are displayed at the top with progress bars '
    'and color-coded status (green for on-target, red for off-target):'
)

add_styled_table(
    ['KPI', 'Target', 'Description'],
    [
        ['Loss Ratio', '65%', 'Claims paid as percentage of premiums earned'],
        ['Expense Ratio', '35%', 'Operating costs as percentage of premium income'],
        ['Combined Ratio', '100%', 'Sum of loss and expense ratios — overall efficiency'],
        ['Retention Rate', '85%', 'Percentage of policies renewed at expiry'],
        ['Collection Rate', '90%', 'Premium collection efficiency against invoiced amount'],
        ['Avg Cycle Time', '30 days', 'Average time from claim filing to settlement'],
        ['Avg Severity', '50,000', 'Average claim payout amount in INR'],
    ],
    col_widths=[1.5, 1.0, 4.0]
)

doc.add_heading('KPI Summary Cards', level=3)
add_bullet_list([
    'Total Policies: 52 (23 active) — with orange document icon',
    'Total Premium: Rs.43,87,475 — with orange currency icon',
    'Total Coverage: Rs.13,73,00,324 — with orange trend icon',
    'Total Claims: 23 (Rs.1,00,64,301 total) — with orange alert icon',
])

doc.add_paragraph()
add_screenshot('03-dashboard-bottom.png', 6.2)

doc.add_heading('Analytics Charts', level=3)
add_bullet_list([
    'Policies by Status — Vertical bar chart showing Active (24), Endorsed, Draft, Lapsed, Cancelled, Expired, Quoted',
    'Policy Type Distribution — Donut chart: Mortgage Guarantee (35%), Credit Protection (33%), Coverage Plus (33%)',
    'Risk Distribution — Pie chart: Low (24), Medium (16), High (10)',
    'Claims by Status — Horizontal bar chart: Approved, Filed, Under Review, Rejected, Settled',
    'Underwriting Decisions — Donut chart: Auto-Approve (28), Refer (12), Reject (10)',
])

doc.add_heading('Dashboard Widgets', level=3)
add_bullet_list([
    'My Tasks — Shows assigned tasks with SLA deadlines and "Open" status badge',
    'Underwriting Overview — Total Evaluations (50) and Average Risk Score (41.5)',
    'Billing Overview — Total Accounts (10), Active (7), Outstanding (Rs.13,058), Collected (Rs.10,033)',
    'Overdue Accounts — Alert widget showing 5 overdue accounts with total overdue Rs.2,950',
])

doc.add_page_break()

# ============================================================================
# 4. POLICY MANAGEMENT
# ============================================================================
doc.add_heading('4. Policy Management', level=1)

add_screenshot('04-policies.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Policy Management module is the core of the PAS platform. It provides '
    'comprehensive policy listing, creation, filtering, and lifecycle management. '
    'Users can create new policies, generate quotes, and manage the entire policy '
    'lifecycle from a single interface.'
)

doc.add_heading('Policy List Features', level=3)
add_bullet_list([
    'Comprehensive data table with 8 columns: Policy ID, Type, Status, Premium, Coverage, Risk, Start Date, Version',
    'Advanced filtering by 12 statuses: Draft, Quoted, Bound, Issued, Active, Endorsed, Renewal Pending, Reinstated, Lapsed, Cancelled, Expired',
    'Filter by policy type: Mortgage Guarantee, Credit Protection, Coverage Plus',
    'Real-time search across all fields with instant results',
    'Sortable columns — click any header to sort ascending/descending',
    'Pagination with 10 records per page and full navigation controls (first, prev, next, last)',
    'Record count display: "Showing 1-10 of 20 records"',
])

doc.add_heading('Quick Actions', level=3)
add_bullet_list([
    'New Quote — Opens the multi-step quoting wizard (Customer Select -> Coverage Config -> Premium Review -> Bind)',
    'New Policy — Creates a new policy directly with full form',
])

doc.add_heading('Policy Status Badges', level=3)
add_styled_table(
    ['Status', 'Badge Color', 'Description'],
    [
        ['Active', 'Green', 'Policy is in force, coverage is active'],
        ['Endorsed', 'Blue', 'Policy has been modified mid-term (endorsement applied)'],
        ['Draft', 'Gray', 'Initial creation, not yet submitted for quoting'],
        ['Quoted', 'Orange', 'Premium has been calculated, awaiting binding'],
        ['Bound', 'Indigo', 'Quote accepted, awaiting issuance'],
        ['Issued', 'Teal', 'Policy formally issued to policyholder'],
        ['Renewal Pending', 'Yellow', 'Renewal initiated, awaiting underwriting/payment'],
        ['Reinstated', 'Cyan', 'Previously lapsed/cancelled policy restored'],
        ['Lapsed', 'Red', 'Premium overdue beyond grace period, coverage suspended'],
        ['Cancelled', 'Red', 'Terminated by request or non-payment'],
        ['Expired', 'Gray', 'Policy term ended, not renewed'],
    ],
    col_widths=[1.5, 1.2, 3.8]
)

doc.add_heading('Risk Level Indicators', level=3)
add_styled_table(
    ['Risk Level', 'Color', 'Indicator'],
    [
        ['Low', 'Green', 'Down arrow — low risk profile'],
        ['Medium', 'Yellow', 'Right arrow — moderate risk profile'],
        ['High', 'Red', 'Up arrow — high risk profile'],
    ],
    col_widths=[1.5, 1.0, 4.0]
)

doc.add_page_break()

# ============================================================================
# 5. POLICY DETAIL & LIFECYCLE
# ============================================================================
doc.add_heading('5. Policy Detail & Lifecycle', level=1)

add_screenshot('05-policy-detail.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Policy Detail page provides a complete view of an individual policy with '
    'a visual lifecycle stepper, action buttons, and an 8-tab detail interface. '
    'This is the central page for all policy operations including endorsements, '
    'renewals, billing, documents, and audit trail.'
)

doc.add_heading('Policy Lifecycle Stepper', level=3)
doc.add_paragraph(
    'The lifecycle stepper displays the complete journey of a policy through five stages:'
)
add_bullet_list([
    'Draft — Initial policy creation (gray circle)',
    'Quoted — Premium calculation completed (green checkmark when passed)',
    'Bound — Quote accepted by customer (green checkmark when passed)',
    'Issued — Policy formally issued (green checkmark when passed)',
    'Active — Policy in force (orange highlighted circle with step number)',
])

doc.add_paragraph(
    'Each completed stage shows a green checkmark. The current stage is highlighted '
    'with an orange circle. Arrows connect the stages to show the flow progression.'
)

doc.add_heading('Available Actions', level=3)
doc.add_paragraph(
    'Based on the current policy status, available transition actions are displayed as '
    'clickable orange/red badges below the stepper. For an Active policy:'
)
add_bullet_list([
    'Endorsed — Apply a mid-term modification to the policy',
    'Renewal Pending — Initiate the renewal process',
    'Lapsed — Mark the policy as lapsed due to non-payment',
    'Cancelled — Cancel the policy by request',
    'Expired — Mark the policy as expired (term ended)',
])

doc.add_heading('Header Action Buttons', level=3)
add_bullet_list([
    'Initiate Renewal — Start the renewal workflow for this policy',
    'Lapse — Lapse the policy (red button, requires confirmation)',
    'Cancel — Cancel the policy (red button, requires confirmation)',
])

doc.add_heading('8-Tab Detail Interface', level=3)
add_styled_table(
    ['Tab', 'Content', 'Details'],
    [
        ['Details', 'Policy & Financial Information', 'Policy ID, Type, Status, Risk Category, Version, Premium, Coverage, Dates, Customer ID, Created timestamp'],
        ['Endorsements', 'Mid-term policy modifications', 'List of endorsements with effective date, type of change, premium delta, and approval status'],
        ['Renewals', 'Renewal history & tracking', 'Previous renewals, upcoming renewal date, renewal underwriting status, premium changes'],
        ['Timeline', 'Chronological event log', 'Visual timeline of all lifecycle events: creation, status changes, endorsements, payments, documents'],
        ['Versions', 'Policy version history', 'Version number, effective date, changes made, comparison between versions'],
        ['Billing', 'Payment information', 'Linked billing account, installment schedule, payment history, outstanding balance'],
        ['Documents', 'Policy documents', 'Generated documents (declarations, certificates), uploaded documents, download/preview actions'],
        ['Audit', 'Complete audit trail', 'Every user action logged: who, what, when — with before/after values for changes'],
    ],
    col_widths=[1.2, 2.0, 3.3]
)

doc.add_page_break()

# ============================================================================
# 6. CUSTOMER MANAGEMENT
# ============================================================================
doc.add_heading('6. Customer Management', level=1)

add_screenshot('06-customers.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Customer Management module provides a centralized view of all customers '
    'with their associated policies, risk profiles, and contact information. It supports '
    'full CRUD operations and real-time search.'
)

doc.add_heading('Customer Table', level=3)
add_bullet_list([
    'Customer ID — Unique identifier (e.g., CUST001)',
    'Name — Full name of the policyholder',
    'Email — Contact email address',
    'Phone — Contact phone number',
    'Status — Active (green badge) or Inactive (gray badge)',
    'Policies — Number of linked policies with clickable navigation',
    'Risk Profile — Low (green), Medium (yellow), High (red) risk assessment',
])

doc.add_heading('Customer Operations', level=3)
add_bullet_list([
    'New Customer — Create a new customer with full profile details',
    'Edit — Modify existing customer information',
    'View Details — Navigate to customer detail page with all linked entities',
    'Search — Real-time search by name, email, or phone number',
    'Pagination — Navigate through customer records with page controls',
])

doc.add_page_break()

# ============================================================================
# 7. UNDERWRITING ENGINE
# ============================================================================
doc.add_heading('7. Underwriting Engine', level=1)

add_screenshot('07-underwriting.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Underwriting Engine provides automated risk assessment with a configurable '
    'rules engine. It evaluates policies against predefined criteria, generates risk '
    'scores (1-100), and makes automated decisions (Auto-Approve, Refer, Reject). '
    'Senior underwriters can manually override automated decisions.'
)

doc.add_heading('Risk Scoring System', level=3)
doc.add_paragraph(
    'Each policy is evaluated against 6 underwriting rules, resulting in a composite '
    'risk score from 1 (lowest risk) to 100 (highest risk):'
)

add_styled_table(
    ['Rule', 'Criteria', 'Impact'],
    [
        ['Property Value Limits', 'Property value within acceptable range', 'Rejects if outside bounds'],
        ['LTV Ratio', 'Loan-to-Value ratio thresholds', 'Higher LTV increases risk score'],
        ['Credit Score', 'Minimum credit score requirements', 'Below threshold triggers refer/reject'],
        ['Employment Verification', 'Employment status and history', 'Unverified employment adds risk'],
        ['Property Type Eligibility', 'Eligible property categories', 'Ineligible types auto-rejected'],
        ['Coverage Ratio', 'Coverage amount relative to value', 'Excessive coverage flagged'],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

doc.add_heading('Decision Categories', level=3)
add_styled_table(
    ['Decision', 'Risk Score Range', 'Action'],
    [
        ['Auto-Approve', '1-35 (Low Risk)', 'Policy automatically approved for binding'],
        ['Refer', '36-65 (Medium Risk)', 'Sent to senior underwriter queue for review'],
        ['Reject', '66-100 (High Risk)', 'Policy automatically declined with reason codes'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading('Advanced Features', level=3)
add_bullet_list([
    'Rule Configuration (Admin) — Create, edit, and manage underwriting rules with IF/THEN logic',
    'Referral Queue — Separate queue for cases requiring senior underwriter review with escalation indicators',
    'Manual Override — Senior underwriters can override automated decisions with mandatory justification',
    'Delegated Authority — Authority limits per underwriter (maximum coverage amount, maximum risk score)',
    'Evaluate Button — One-click evaluation with instant risk score calculation and decision',
    'Risk Scorecard — Visual breakdown by category: Financial, Property, Customer, Compliance',
])

doc.add_page_break()

# ============================================================================
# 8. CLAIMS MANAGEMENT
# ============================================================================
doc.add_heading('8. Claims Management', level=1)

add_screenshot('08-claims.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Claims Management module handles the complete claims lifecycle from '
    'First Notice of Loss (FNOL) intake through investigation, evaluation, and '
    'settlement. It includes reserve management, fraud scoring, and loss mitigation tracking.'
)

doc.add_heading('Claims Table', level=3)
add_bullet_list([
    'Claim ID — Unique identifier (e.g., CLM-2025-00001)',
    'Policy — Linked policy number for cross-reference',
    'Type — Claim category (Property Damage, Income Loss, etc.)',
    'Status — Current claim status with color-coded badge',
    'Amount — Claimed amount in INR',
    'Filed Date — Date of FNOL submission',
])

doc.add_heading('Claims Status Workflow', level=3)
doc.add_paragraph('Claims progress through the following status flow:')

add_styled_table(
    ['Status', 'Description', 'Next Actions'],
    [
        ['Filed', 'FNOL submitted, claim registered', 'Assign to adjuster, begin investigation'],
        ['Under Review', 'Investigation in progress', 'Evaluate evidence, assess damage'],
        ['Approved', 'Claim validated and approved', 'Set reserves, initiate settlement'],
        ['Rejected', 'Claim denied with reason codes', 'Send denial letter, close claim'],
        ['Settled', 'Payment issued to claimant', 'Archive documentation, close claim'],
    ],
    col_widths=[1.3, 2.5, 2.7]
)

doc.add_heading('FNOL Intake Process', level=3)
doc.add_paragraph(
    'The FNOL (First Notice of Loss) wizard guides users through a structured 5-step intake process:'
)

add_styled_table(
    ['Step', 'Name', 'Data Collected'],
    [
        ['1', 'Incident Information', 'Date, time, location, description of incident'],
        ['2', 'Policy Lookup', 'Search and link to the insured policy'],
        ['3', 'Damage Details', 'Type of damage, extent, estimated amount, photographs'],
        ['4', 'Document Upload', 'Supporting documentation (police reports, medical records, invoices)'],
        ['5', 'Review & Submit', 'Final review of all information before formal filing'],
    ],
    col_widths=[0.5, 1.5, 4.5]
)

doc.add_heading('Additional Features', level=3)
add_bullet_list([
    'Reserve Management — Set initial reserves, make adjustments, view reserve history chart',
    'Fraud Indicator Scoring — Rule-based scoring for duplicate claims, timing patterns, amount thresholds',
    'Claims Timeline — Visual chronological log of all claim events (FNOL, notes, reserve changes, payments)',
    'Loss Mitigation — Track workout plans, forbearance status, and loan modification details',
])

doc.add_page_break()

# ============================================================================
# 9. BILLING & PAYMENTS
# ============================================================================
doc.add_heading('9. Billing & Payments', level=1)

add_screenshot('09-billing.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Billing & Payments module provides a complete financial lifecycle including '
    'invoicing, payment tracking, installment plans, receivables aging, and double-entry '
    'ledger accounting. It integrates directly with the policy lifecycle for automated '
    'billing and grace period management.'
)

doc.add_heading('Billing Account Management', level=3)
add_bullet_list([
    'Billing accounts are automatically linked to policies at issuance',
    'Each account tracks balance, payment plan, installment schedule, and payment history',
    'Account statuses: Active, Suspended (grace period), Closed',
    'Quick-view KPI cards: Total Accounts, Active Accounts, Outstanding Amount, Collected Amount',
])

doc.add_heading('Invoicing', level=3)
add_bullet_list([
    'Auto-generated invoices with detailed line items (premium, fees, taxes)',
    'Invoice statuses: Pending, Paid, Overdue, Void',
    'Due date tracking with overdue alerts',
    'Print and download invoice capability',
])

doc.add_heading('Payment Tracking', level=3)
add_styled_table(
    ['Payment Method', 'Details'],
    [
        ['Cash', 'Cash payments with receipt reference'],
        ['Check', 'Check number and bank details'],
        ['UPI', 'UPI transaction ID'],
        ['Bank Transfer', 'NEFT/RTGS reference number'],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_heading('Installment Plans', level=3)
add_bullet_list([
    'Annual — Single payment for the full premium',
    'Semi-Annual — Two equal installments, 6 months apart',
    'Quarterly — Four equal installments, 3 months apart',
    'Monthly — Twelve equal installments with monthly due dates',
])

doc.add_heading('Receivables Aging', level=3)
doc.add_paragraph(
    'A visual aging chart displays outstanding receivables categorized by age:'
)
add_bullet_list([
    'Current — Invoices within payment terms',
    '30-day — 1-30 days past due',
    '60-day — 31-60 days past due',
    '90-day — 61-90 days past due',
    '90+ day — More than 90 days past due (critical)',
])

doc.add_heading('Grace Period Automation', level=3)
doc.add_paragraph(
    'The system automates the grace period workflow: when an invoice becomes overdue, '
    'a grace period timer begins. If payment is not received within the grace period, '
    'the linked policy is automatically flagged for lapse processing.'
)

doc.add_page_break()

# ============================================================================
# 10. TASK MANAGEMENT
# ============================================================================
doc.add_heading('10. Task Management', level=1)

add_screenshot('10-tasks.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Task Management module provides a centralized work queue with SLA tracking, '
    'priority management, and automatic task generation from workflow events. It ensures '
    'nothing falls through the cracks and all team members have clear work assignments.'
)

doc.add_heading('Task Queue Features', level=3)
add_bullet_list([
    'Summary Cards — At-a-glance counts for Open, In Progress, Completed, and Overdue tasks',
    'Task Table — Title, Assignee, Priority, Due Date, Status columns with sorting',
    'Priority Levels — Urgent (red), High (orange), Medium (yellow), Low (green)',
    'SLA Tracking — Due date countdown timers with overdue indicators and alerts',
    'Assignment — Assign or reassign tasks to any team member',
    'Task Detail Drawer — Side panel with full task details, activity log, and actions',
])

doc.add_heading('Workflow-Driven Auto-Assignment', level=3)
add_styled_table(
    ['Trigger Event', 'Auto-Created Task', 'Assigned To'],
    [
        ['New claim filed', 'Review and investigate claim', 'Claims Adjuster'],
        ['UW referral triggered', 'Review referred application', 'Senior Underwriter'],
        ['Policy renewal due', 'Process policy renewal', 'Operations Team'],
        ['Payment overdue', 'Follow up on overdue payment', 'Billing Team'],
        ['Document required', 'Collect required documentation', 'Agent'],
        ['Compliance deadline', 'Complete regulatory filing', 'Compliance Officer'],
    ],
    col_widths=[2.0, 2.5, 2.0]
)

doc.add_page_break()

# ============================================================================
# 11. DOCUMENT MANAGEMENT
# ============================================================================
doc.add_heading('11. Document Management', level=1)

add_screenshot('11-documents.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Document Management module provides a centralized repository for all policy, '
    'claims, underwriting, and correspondence documents. It supports file upload, '
    'template-based document generation, categorization, and version control.'
)

doc.add_heading('Document Categories', level=3)
add_styled_table(
    ['Category', 'Document Types', 'Description'],
    [
        ['Policy Documents', 'Declarations, Certificates, Endorsement Schedules', 'Official policy documentation issued to policyholders'],
        ['Claims Documents', 'FNOL Reports, Investigation Notes, Settlement Letters', 'Documentation related to claims processing'],
        ['UW Documents', 'Risk Assessments, Rule Evaluations, Referral Notes', 'Underwriting analysis and decision records'],
        ['Correspondence', 'Customer Letters, Renewal Notices, Premium Receipts', 'All customer-facing communications'],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_heading('Key Features', level=3)
add_bullet_list([
    'Drag-and-Drop Upload — File upload with progress bar and size display',
    'Document Generation — Auto-generate documents from templates (declarations, certificates, invoices)',
    'In-App Preview — View PDFs and images directly within the application',
    'Full-Text Search — Search across document names and metadata',
    'Version Control — Track document versions with history',
    'Linked Entities — Documents are linked to their parent policy, claim, or customer',
    'Bulk Download — Select multiple documents for ZIP download',
])

doc.add_page_break()

# ============================================================================
# 12. REPORTS & ANALYTICS
# ============================================================================
doc.add_heading('12. Reports & Analytics', level=1)

add_screenshot('12-reports.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Reports module provides comprehensive analytics across policies, claims, '
    'and financial metrics. It features interactive charts, date range filtering, '
    'and summary metric cards.'
)

doc.add_heading('Report Types', level=3)
add_styled_table(
    ['Tab', 'Reports', 'Metrics'],
    [
        ['Policy Reports', 'Premium trends, status distribution, type breakdown', 'Active policies, total premium, avg premium, new policies'],
        ['Claims Reports', 'Claims volume, severity trends, cycle time analysis', 'Total claims, total paid, avg claim amount, open claims'],
        ['Financial Reports', 'Revenue tracking, loss ratios, expense analysis', 'Revenue, expenses, net income, loss ratio trend'],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_heading('Date Range Filtering', level=3)
add_bullet_list([
    'This Month — Current calendar month data',
    'Last Month — Previous calendar month data',
    'Last Quarter — Previous 3-month quarter',
    'Year-to-Date (YTD) — January 1 through today',
    'Custom Range — Select specific start and end dates',
])

doc.add_heading('Chart Types', level=3)
add_bullet_list([
    'Vertical Bar Charts — Policy counts, premium trends (orange bars)',
    'Pie / Donut Charts — Distribution charts for types, statuses, risk levels',
    'Horizontal Bar Charts — Claims by status comparison',
    'Summary Cards — Key metrics with trend indicators',
])

doc.add_page_break()

# ============================================================================
# 13. ADMINISTRATION
# ============================================================================
doc.add_heading('13. Administration & User Management', level=1)

add_screenshot('13-admin.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Administration module provides user management, role-based access control, '
    'and system-wide audit logging. Only Admin-role users have access to this module.'
)

doc.add_heading('User Management', level=3)
add_bullet_list([
    'User Table — Name, Email, Role, Status, Last Active, Actions columns',
    'Add User — Create new user accounts with role assignment',
    'Edit User — Modify user profile, change roles, update permissions',
    'Activate/Deactivate — Enable or disable user accounts',
    'Color-coded Role Badges — Visual role identification at a glance',
    'Search — Find users by name, email, or role',
])

doc.add_heading('Audit Logging', level=3)
doc.add_paragraph(
    'Every administrative action is logged with: User who performed the action, '
    'timestamp, action type, affected entity, and before/after values for data changes. '
    'The audit log is immutable and accessible only to Admin and Auditor roles.'
)

doc.add_page_break()

# ============================================================================
# 14. COMPLIANCE MANAGEMENT
# ============================================================================
doc.add_heading('14. Compliance Management', level=1)

add_screenshot('14-compliance.png', 6.2)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The Compliance Management module tracks regulatory requirements, filing deadlines, '
    'and audit findings. It ensures the organization maintains compliance with IRDAI '
    'regulations and internal governance standards.'
)

doc.add_heading('Compliance Tracking', level=3)
add_styled_table(
    ['Status', 'Badge Color', 'Meaning'],
    [
        ['Compliant', 'Green', 'Requirement is fully met with documentation'],
        ['Non-Compliant', 'Red', 'Requirement not met — action required'],
        ['Under Review', 'Yellow', 'Compliance review in progress'],
        ['Due', 'Orange', 'Upcoming deadline — attention needed'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

doc.add_heading('Features', level=3)
add_bullet_list([
    'Regulatory requirements tracker with status indicators',
    'Compliance calendar with upcoming regulatory deadlines',
    'Alert system for approaching and missed deadlines',
    'IRDAI regulatory filing status and submission tracking',
    'SOC 2 compliance documentation management',
    'Compliance report generation for internal and external audits',
    'Complete audit trail for all compliance-related actions',
])

doc.add_page_break()

# ============================================================================
# 15. TECHNOLOGY STACK
# ============================================================================
doc.add_heading('15. Technology Stack & Architecture', level=1)

doc.add_heading('Application Architecture', level=2)
add_styled_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend Framework', 'Next.js 14, React 18', 'Server-side rendering, component architecture, routing'],
        ['Language', 'TypeScript', 'Type safety across frontend and backend'],
        ['Styling', 'Tailwind CSS', 'Utility-first CSS with enterprise design system'],
        ['Icons', 'Lucide React', 'Consistent icon set across the application'],
        ['Charts', 'Recharts', 'Interactive data visualizations (bar, pie, donut, line)'],
        ['Backend', 'Express.js', 'RESTful API server with middleware architecture'],
        ['Authentication', 'JWT + Refresh Tokens', 'Stateless, scalable authentication'],
        ['Authorization', 'RBAC Middleware', 'Role-based permission checking on every request'],
        ['Storage', 'JSON File Persistence', 'Rapid prototyping (database-ready architecture)'],
    ],
    col_widths=[1.8, 2.0, 2.7]
)

doc.add_heading('Non-Functional Features', level=2)
add_styled_table(
    ['Feature', 'Implementation'],
    [
        ['Responsive Design', 'Works across desktop (1440px+), tablet (768px), and mobile (375px)'],
        ['Accessibility', 'WCAG 2.1 compliant — skip navigation, ARIA labels, keyboard navigation, focus management'],
        ['Performance', 'Skeleton loading states, staggered animations, optimistic updates, lazy loading'],
        ['Security', '256-bit encryption, JWT tokens, bcrypt password hashing, CORS, input sanitization'],
        ['UX Patterns', 'Toast notifications, confirmation dialogs, breadcrumb navigation, global search (Cmd+K)'],
        ['Error Handling', 'Graceful error boundaries, user-friendly error messages, retry mechanisms'],
        ['Theme', 'Light/dark mode support, default light theme, 30% orange accent coverage'],
    ],
    col_widths=[1.8, 4.7]
)

doc.add_page_break()

# ============================================================================
# 16. USER ROLES & PERMISSIONS
# ============================================================================
doc.add_heading('16. User Roles & Permissions', level=1)

doc.add_paragraph(
    'The PAS platform implements Role-Based Access Control (RBAC) with 5 predefined roles. '
    'Each role has specific permissions that control access to modules and actions.'
)

add_styled_table(
    ['Module', 'Admin', 'Underwriter', 'Claims Adj.', 'Agent', 'Auditor'],
    [
        ['Dashboard', 'Full', 'Full', 'Full', 'Full', 'Read'],
        ['Policies', 'Full CRUD', 'Read', 'Read', 'Create/Read', 'Read'],
        ['Customers', 'Full CRUD', 'Read', 'Read', 'Full CRUD', 'Read'],
        ['Underwriting', 'Full', 'Full', 'Read', 'Read', 'Read'],
        ['Referrals', 'Full', 'Full', 'None', 'None', 'Read'],
        ['Claims', 'Full CRUD', 'Read', 'Full CRUD', 'Create', 'Read'],
        ['FNOL', 'Full', 'None', 'Full', 'Create', 'Read'],
        ['Billing', 'Full', 'Read', 'Read', 'Read', 'Read'],
        ['Tasks', 'Full', 'Assigned', 'Assigned', 'Assigned', 'Read'],
        ['Documents', 'Full', 'Read', 'Read', 'Upload', 'Read'],
        ['Reports', 'Full', 'Full', 'Full', 'Limited', 'Full'],
        ['Compliance', 'Full', 'None', 'None', 'None', 'Read'],
        ['Admin', 'Full', 'None', 'None', 'None', 'None'],
    ],
    col_widths=[1.3, 1.0, 1.0, 1.0, 1.0, 1.0]
)

doc.add_page_break()

# ============================================================================
# 17. NAVIGATION REFERENCE
# ============================================================================
doc.add_heading('17. Navigation Reference', level=1)

doc.add_paragraph(
    'The sidebar navigation is organized into 4 sections with role-based visibility:'
)

add_styled_table(
    ['Section', 'Module', 'Route', 'Access'],
    [
        ['Core Operations', 'Dashboard', '/dashboard', 'All roles'],
        ['Core Operations', 'Policies', '/policies', 'All roles'],
        ['Core Operations', 'Renewals', '/renewals', 'All roles'],
        ['Core Operations', 'Customers', '/customers', 'All roles'],
        ['Core Operations', 'Underwriting', '/underwriting', 'Admin, UW, Ops, Viewer'],
        ['Core Operations', 'Claims', '/claims', 'Admin, Claims, Ops, Viewer'],
        ['Core Operations', 'FNOL', '/claims/fnol', 'Admin, Claims, Ops'],
        ['Financial', 'Billing', '/billing', 'Admin, Ops, Viewer'],
        ['Financial', 'Tasks', '/tasks', 'Admin, UW, Claims, Ops'],
        ['Reports & Docs', 'Documents', '/documents', 'Admin, Ops'],
        ['Reports & Docs', 'Reports', '/reports', 'All roles'],
        ['Reports & Docs', 'Docs', '/documentation', 'All roles'],
        ['Administration', 'Compliance', '/admin/compliance', 'Admin only'],
        ['Administration', 'Admin', '/admin', 'Admin only'],
    ],
    col_widths=[1.5, 1.3, 2.0, 1.7]
)

doc.add_page_break()

# ============================================================================
# 18. API ENDPOINTS REFERENCE
# ============================================================================
doc.add_heading('18. API Endpoints Reference', level=1)

doc.add_paragraph(
    'All API endpoints are prefixed with /api/v1/ and require JWT authentication (except /auth/login).'
)

add_styled_table(
    ['Module', 'Base Route', 'Key Endpoints'],
    [
        ['Authentication', '/api/v1/auth', 'POST /login, POST /refresh, POST /signup, GET /profile'],
        ['Policies', '/api/v1/policies', 'GET /, POST /, GET /:id, PUT /:id, POST /:id/endorse, POST /:id/renew'],
        ['Customers', '/api/v1/customers', 'GET /, POST /, GET /:id, PUT /:id, DELETE /:id'],
        ['Underwriting', '/api/v1/underwriting', 'GET /, POST /evaluate, GET /:id, PUT /:id/override'],
        ['UW Rules', '/api/v1/uw', 'GET /rules, POST /rules, PUT /rules/:id, DELETE /rules/:id'],
        ['Claims', '/api/v1/claims', 'GET /, POST /, GET /:id, PUT /:id, POST /fnol, PUT /:id/reserve'],
        ['Billing', '/api/v1/billing', 'GET /accounts, GET /:id, POST /invoices, POST /payments'],
        ['Tasks', '/api/v1/tasks', 'GET /, POST /, GET /:id, PUT /:id, PUT /:id/assign'],
        ['Documents', '/api/v1/documents', 'GET /, POST /upload, GET /:id/download, POST /generate'],
        ['Reports', '/api/v1/reports', 'GET /policy, GET /claims, GET /executive, GET /financial'],
        ['Dashboard', '/api/v1/dashboard', 'GET /, GET /:role'],
        ['Notifications', '/api/v1/notifications', 'GET /, PUT /:id/read, PUT /read-all'],
        ['Admin', '/api/v1/admin', 'GET /users, POST /users, PUT /users/:id, GET /audit-logs'],
        ['Compliance', '/api/v1/compliance', 'GET /requirements, PUT /requirements/:id'],
        ['Activity', '/api/v1/activity', 'GET / (paginated activity feed)'],
        ['Products', '/api/v1/products', 'GET /, POST /, PUT /:id'],
        ['Bulk Ops', '/api/v1/bulk', 'POST /:operation, GET /:operationId/status'],
        ['Integrations', '/api/v1/integrations', 'GET /status, POST /webhooks'],
    ],
    col_widths=[1.2, 1.8, 3.5]
)

doc.add_paragraph()
add_note_box(
    'All endpoints return JSON responses. Authentication uses Bearer tokens in the '
    'Authorization header. Pagination is supported with ?page= and ?limit= query parameters.',
    prefix='API Note'
)

doc.add_page_break()

# ============================================================================
# FINAL PAGE
# ============================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('IMGC Corporation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Policy Administration System — Application Guide')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 3.2.1  |  Enterprise Edition  |  March 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Confidential — For authorized recipients only')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('India Mortgage Guarantee Corporation')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ============================================================================
# SAVE
# ============================================================================
doc.save(OUTPUT_FILE)
print(f'Application Guide saved to: {OUTPUT_FILE}')
print(f'Total sections: 18')
