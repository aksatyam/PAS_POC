"""
Generates APPLICATION_FLOW_DESIGN.docx
IMGC PAS — Application Flow Design Document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(table, color="CCCCCC"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color)
                tcB.append(b)
            tcPr.append(tcB)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    run.font.size = Pt(16)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    run.font.size = Pt(13)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.size = Pt(11)
    return p

def para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    r = p.add_run(text); r.font.size = Pt(10)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    set_para_bg(p, "F4F4F4")
    return p

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def table(doc, headers, rows, hbg="1B3A6B", hfg="FFFFFF", alt="EBF0FA"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        set_cell_bg(c, hbg)
        rn = c.paragraphs[0].runs[0]
        rn.bold = True; rn.font.size = Pt(9)
        rn.font.color.rgb = RGBColor.from_string(hfg)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        bg = alt if ri % 2 == 0 else "FFFFFF"
        r = t.rows[ri+1]
        for ci, cell_text in enumerate(row):
            c = r.cells[ci]
            c.text = str(cell_text)
            set_cell_bg(c, bg)
            c.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_border(t)
    return t

def callout(doc, text, bg="E8F4FD", border="2E75B6"):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, bg)
    p = c.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10); r.italic = True
    set_cell_border(t, border)
    doc.add_paragraph()


# ── build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Cover ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("IMGC PAS\nApplication Flow Design Document")
    tr.bold = True; tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    doc.add_paragraph()
    for lbl, val in [
        ("System:", "India Mortgage Guarantee Corporation — Policy Administration System"),
        ("Document Classification:", "Internal — Architecture & Engineering"),
        ("Version:", "1.0"),
        ("Date:", "March 2026"),
        ("Audience:", "Enterprise Architects, Senior Engineers, Technical Leadership"),
    ]:
        mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = mp.add_run(lbl + "  "); r1.bold = True; r1.font.size = Pt(10)
        r2 = mp.add_run(val); r2.font.size = Pt(10)
    doc.add_page_break()

    # ── TOC ────────────────────────────────────────────────────────────────
    h1(doc, "Table of Contents")
    for item in [
        "1. High-Level Design (HLD)",
        "2. Low-Level Design (LLD)",
        "3. Entity-Relationship Diagram",
        "4. State Machine Diagrams",
        "5. Sequence Diagrams",
        "6. Data Flow Diagrams",
        "7. API Design Reference",
        "8. Deployment Architecture",
    ]:
        bullet(doc, item)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1 — HLD
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "1. High-Level Design (HLD)")

    h2(doc, "1.1 System Context")
    para(doc,
        "The IMGC PAS is an enterprise Policy Administration System for mortgage guarantee and credit "
        "protection operations. It serves multiple actor types — lenders, borrowers, operations staff, "
        "underwriters, claims handlers, auditors, and regulators — through a secured, role-governed "
        "web application backed by a RESTful API.")
    doc.add_paragraph()
    callout(doc,
        "The system manages the complete insurance lifecycle: Quote → Bind → Issue → Endorse → Renew → "
        "Cancel / Expire, alongside Claims processing (FNOL → Adjudication → Settlement) and Billing "
        "(Invoice → Payment → Ledger).")

    h2(doc, "1.2 Architectural Layers")
    table(doc,
        ["Layer", "Technology", "Responsibility"],
        [
            ["Presentation", "Next.js 14, React 18, TypeScript, Tailwind CSS", "20+ modules, 45+ reusable components, role-based dashboards"],
            ["State Management", "Zustand + TanStack Query", "Client-side state, server-state cache, optimistic updates"],
            ["API Gateway", "NGINX + WAF + Rate Limiter", "TLS termination, WAF rules, per-IP/user rate limiting"],
            ["Application", "Express.js + TypeScript", "18 route files, 18 controllers, 19 services — full business logic"],
            ["Security Middleware", "JWT + Zod + RBAC", "Token verification, schema validation, role enforcement"],
            ["Data", "PostgreSQL + Redis", "Relational persistence, session cache, job queues"],
            ["File Store", "S3 / Azure Blob", "Policy documents, claim evidence, generated certificates"],
            ["Audit", "Immutable Audit DB", "Tamper-proof activity trail for all state mutations"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "1.3 Module Catalogue")
    table(doc,
        ["Module", "Responsibility", "Key Entities"],
        [
            ["Auth / RBAC",     "Login, JWT issuance, role enforcement",               "User, Role, Session"],
            ["Policy",          "Full policy lifecycle management",                     "Policy, PolicyVersion, Endorsement, Renewal"],
            ["QDE",             "Quick Data Entry — initial quote capture",             "Draft Policy, Customer Ref"],
            ["DDE",             "Detailed Data Entry — complete application",           "Policy, Documents, UW trigger"],
            ["Underwriting",    "Risk evaluation, rule engine, referrals",             "UWRecord, Rule, Referral, Authority"],
            ["Claims",          "FNOL, adjudication, reserves, fraud scoring",         "Claim, FNOL, Reserve, FraudAssessment"],
            ["Billing",         "Invoicing, payments, ledger, aging",                  "BillingAccount, Invoice, Payment, Ledger"],
            ["Servicing",       "NPA, delinquency, premium checks",                    "NPA Record, Delinquency"],
            ["Documents",       "Upload, versioning, template generation",             "PolicyDocument, Template"],
            ["Tasks",           "Work item assignment and tracking",                   "Task, Assignment"],
            ["Notifications",   "Email/SMS dispatch",                                  "Notification, Template"],
            ["Reports",         "Analytics, KPI dashboards, exports",                  "Report, KPI Metric"],
            ["Compliance",      "Regulatory obligation tracking",                       "ComplianceRequirement"],
            ["Admin",           "User, role, product, master data setup",              "User, Product, Config"],
            ["Audit",           "Immutable activity trail",                            "AuditLog, ActivityEvent"],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2 — LLD
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "2. Low-Level Design (LLD)")

    h2(doc, "2.1 Request Processing Pipeline")
    steps = [
        ("1. CORS Check",        "Allow-listed origins only; OPTIONS pre-flight handled"),
        ("2. Rate Limiter",      "100 req / 15 min per IP; 1,000 req / min per authenticated user"),
        ("3. JWT Auth",          "Verify Bearer token signature and expiry; extract userId, email, role"),
        ("4. RBAC Guard",        "Match role against route permission matrix; 403 if insufficient"),
        ("5. Zod Validation",    "Parse and validate request body / query params against schema"),
        ("6. Controller",        "Extract params, call service, format standard response envelope"),
        ("7. Service Layer",     "All business logic, orchestration, cross-service calls"),
        ("8. Repository / DAO",  "Data read/write, Redis cache check/set, parameterised queries"),
        ("9. Audit Logger",      "Async write to immutable audit_logs (userId, action, before/after)"),
    ]
    for step, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"{step}: "); r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(desc); r2.font.size = Pt(10)

    doc.add_paragraph()
    h2(doc, "2.2 RBAC Permission Matrix")
    table(doc,
        ["Action", "Admin", "Operations", "Underwriter", "Claims", "Viewer"],
        [
            ["User Management",         "✅", "❌", "❌", "❌", "❌"],
            ["Policy — Create (QDE/DDE)","✅", "✅", "❌", "❌", "❌"],
            ["Policy — View",            "✅", "✅", "✅", "✅", "✅"],
            ["Policy — Endorse / Renew", "✅", "✅", "❌", "❌", "❌"],
            ["Policy — Cancel",          "✅", "✅", "❌", "❌", "❌"],
            ["UW — Evaluate",            "✅", "❌", "✅", "❌", "❌"],
            ["UW — Override",            "✅", "❌", "✅ (limit)", "❌", "❌"],
            ["Claims — File FNOL",       "✅", "✅", "❌", "✅", "❌"],
            ["Claims — Adjudicate",      "✅", "❌", "❌", "✅", "❌"],
            ["Claims — Settle",          "✅", "❌", "❌", "✅", "❌"],
            ["Billing — Invoice",        "✅", "✅", "❌", "❌", "❌"],
            ["Billing — Record Payment", "✅", "✅", "❌", "❌", "❌"],
            ["Reports — View",           "✅", "✅", "✅", "✅", "✅"],
            ["Audit Logs — View",        "✅", "❌", "❌", "❌", "❌"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "2.3 Backend Directory Structure")
    code_block(doc, """backend/src/
├── app.ts                     # Express app bootstrap
├── server.ts                  # HTTP server entry point
├── config/
│   ├── env.ts                 # Env variable validation (Zod)
│   ├── jwt.ts                 # JWT secrets and options
│   └── cors.ts                # CORS allow-list
├── middleware/
│   ├── auth.middleware.ts      # JWT extraction + verification
│   ├── rbac.middleware.ts      # Role-based access control
│   ├── validate.middleware.ts  # Zod request body validation
│   ├── rateLimiter.ts          # Per-IP / per-user rate limiting
│   ├── audit.middleware.ts     # Immutable audit log writer
│   ├── errorHandler.ts         # Centralised error formatter
│   └── requestLogger.ts        # Structured Winston logging
├── routes/                    # 18 thin route files
├── controllers/               # 18 controllers (HTTP in/out only)
├── services/                  # 19 services (all business logic)
├── repositories/              # DAOs — data access objects
├── models/                    # TypeScript interfaces and types
├── validators/                # Zod schemas per entity
└── utils/                     # jwt, crypto, pagination, response""")

    doc.add_paragraph()
    h2(doc, "2.4 Frontend Module Map")
    code_block(doc, """frontend/app/
├── (auth)/login/              # Public login page
└── (main)/                    # Protected — JWT required
    ├── dashboard/             # Executive KPI dashboard
    ├── service-desk/          # QDE intake form
    ├── dde/                   # Detailed data entry
    ├── policies/[id]/         # Policy overview + endorsements + renewals
    ├── underwriting/          # UW queue + evaluation
    ├── claims/[id]/           # Claims list + adjudication
    ├── fnol/                  # FNOL intake
    ├── billing/               # Billing accounts + invoices
    ├── customers/             # Customer profiles
    ├── documents/             # Document repository
    ├── tasks/                 # Work item queue
    ├── reports/               # Analytics + exports
    ├── compliance/            # Compliance register
    ├── audit-logs/            # Audit trail viewer
    ├── admin/                 # System administration
    └── master-setup/          # Reference data configuration""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3 — ER DIAGRAM (described in table form)
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "3. Entity-Relationship Diagram")
    callout(doc,
        "The full ER diagram is defined in Mermaid erDiagram format in the companion Markdown file "
        "(APPLICATION_FLOW_DESIGN.md). The tables below describe every entity, its key attributes, "
        "and its relationships.")

    entities = [
        ("USERS", ["id (PK)", "email (UK)", "name", "role", "hashedPassword", "isActive"],
         "5 roles: Admin, Operations, Underwriter, Claims, Viewer"),
        ("SESSIONS", ["id (PK)", "userId (FK)", "refreshToken", "expiresAt", "ipAddress"],
         "JWT refresh token store — Redis backed"),
        ("AUDIT_LOGS", ["id (PK)", "userId (FK)", "action", "entityType", "entityId", "before (jsonb)", "after (jsonb)", "ipAddress"],
         "Immutable — no UPDATE / DELETE permitted"),
        ("CUSTOMERS", ["id (PK)", "name", "email (UK)", "phone", "panNumber", "aadhaarNumber", "address", "kycStatus"],
         "Encrypted PAN / Aadhaar at column level"),
        ("PRODUCTS", ["id (PK)", "name", "code (UK)", "policyType", "coverageOptions (jsonb)", "ratingFactors (jsonb)", "eligibilityCriteria (jsonb)"],
         "Master product configuration — versioned"),
        ("POLICIES", ["id (PK)", "customerId (FK)", "productId (FK)", "policyNumber (UK)", "status", "premiumAmount", "coverageAmount", "startDate", "endDate", "version"],
         "Central entity — 11 possible status values"),
        ("POLICY_VERSIONS", ["id (PK)", "policyId (FK)", "version", "changes (jsonb)", "changedBy (FK)"],
         "Full change history per version increment"),
        ("ENDORSEMENTS", ["id (PK)", "policyId (FK)", "type", "changes (jsonb)", "premiumDelta", "effectiveDate", "status"],
         "5 endorsement types; 4 status values"),
        ("RENEWALS", ["id (PK)", "originalPolicyId (FK)", "renewedPolicyId (FK)", "renewalDate", "newPremiumAmount", "status"],
         "Links original policy to renewed policy"),
        ("UW_RULES", ["id (PK)", "name", "category", "field", "operator", "value (jsonb)", "decision", "priority", "isActive"],
         "Rule engine — evaluated in priority order"),
        ("UW_RECORDS", ["id (PK)", "policyId (FK)", "creditScore", "income", "propertyValue", "ltvRatio", "riskScore", "decision", "ruleResults (jsonb)"],
         "One record per UW evaluation"),
        ("UW_REFERRALS", ["id (PK)", "underwritingId (FK)", "policyId (FK)", "assignedTo (FK)", "status", "priority", "slaDeadline"],
         "SLA-tracked referral queue for manual UW review"),
        ("UW_AUTHORITY", ["id (PK)", "userId (FK)", "maxCoverageAmount", "maxRiskScore", "allowedDecisions (jsonb)", "canOverride"],
         "Delegated authority limits per underwriter"),
        ("FNOL", ["id (PK)", "policyId (FK)", "claimType", "incidentDate", "estimatedAmount", "partiesInvolved (jsonb)", "status"],
         "First Notice of Loss — precursor to Claim"),
        ("CLAIMS", ["id (PK)", "policyId (FK)", "fnolId (FK)", "claimNumber (UK)", "status", "adjudicationStatus", "amount", "settlementAmount", "fraudScore"],
         "Central claims entity; links to reserves, fraud, documents"),
        ("RESERVES", ["id (PK)", "claimId (FK)", "type", "amount", "previousAmount", "reason"],
         "Reserve types: Initial, Adjustment, IBNR"),
        ("FRAUD_ASSESSMENTS", ["id (PK)", "claimId (FK)", "score (0-100)", "level", "indicators (jsonb)"],
         "Automated scoring + rule-based indicators"),
        ("LOSS_MITIGATIONS", ["id (PK)", "claimId (FK)", "type", "status", "startDate", "terms"],
         "Workout plans, forbearance, loan modifications"),
        ("BILLING_ACCOUNTS", ["id (PK)", "policyId (FK)", "customerId (FK)", "paymentFrequency", "totalPremium", "balance", "status", "autopay"],
         "One billing account per policy"),
        ("INVOICES", ["id (PK)", "billingAccountId (FK)", "invoiceNumber (UK)", "amount", "amountPaid", "dueDate", "status", "lineItems (jsonb)"],
         "5 status values; supports partial payment"),
        ("PAYMENTS", ["id (PK)", "billingAccountId (FK)", "invoiceId (FK)", "amount", "method", "status", "processedDate"],
         "5 payment methods: ACH, Wire, Check, Card, Escrow"),
        ("LEDGER_ENTRIES", ["id (PK)", "billingAccountId (FK)", "type", "debit", "credit", "balance", "referenceId"],
         "Running balance ledger — append-only"),
        ("INSTALLMENT_PLANS", ["id (PK)", "billingAccountId (FK)", "numberOfInstallments", "installmentAmount", "installments (jsonb)"],
         "Frequency: Annual, Semi-Annual, Quarterly, Monthly"),
        ("DOCUMENTS", ["id (PK)", "policyId (FK)", "claimId (FK)", "type", "category", "filename", "version", "parentDocumentId (FK)"],
         "23 document types; version-chained"),
        ("DOCUMENT_TEMPLATES", ["id (PK)", "name", "type", "fields (jsonb)"],
         "8 template types: Policy Declaration, COI, Endorsement Schedule, etc."),
        ("TASKS", ["id (PK)", "title", "type", "status", "priority", "assignedTo (FK)", "relatedEntityType", "dueDate"],
         "Work item queue across all modules"),
        ("NOTIFICATIONS", ["id (PK)", "userId (FK)", "type", "channel", "isRead", "relatedEntityType", "sentAt"],
         "Email / SMS / in-app; read-receipt tracked"),
        ("COMPLIANCE_REQUIREMENTS", ["id (PK)", "name", "category", "authority", "status", "priority", "recurrence", "dueDate"],
         "GDPR, POPIA, PCI-DSS, SOC2, ISO27001"),
        ("BULK_OPERATIONS", ["id (PK)", "type", "status", "totalItems", "successCount", "failureCount", "results (jsonb)"],
         "Async batch: Bulk Renewal, Cancel, Invoice, Claim Update"),
    ]

    for ent_name, attrs, note in entities:
        h3(doc, ent_name)
        for a in attrs:
            bullet(doc, a)
        if note:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(f"Note: {note}")
            r.italic = True; r.font.size = Pt(9)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4 — STATE MACHINES
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "4. State Machine Diagrams")
    callout(doc,
        "Full Mermaid stateDiagram-v2 definitions are in APPLICATION_FLOW_DESIGN.md. "
        "This section provides the state transition tables for each key entity.")

    h2(doc, "4.1 Policy Lifecycle States")
    table(doc,
        ["From State", "Event / Trigger", "To State", "Actor"],
        [
            ["—",                "QDE form submitted",                 "Draft",           "Operations"],
            ["Draft",            "Premium calculated",                 "Quoted",          "System"],
            ["Draft",            "Withdrawn",                          "Cancelled",       "Operations"],
            ["Quoted",           "Customer accepts",                   "Bound",           "Operations"],
            ["Quoted",           "Terms revised",                      "Draft",           "Operations"],
            ["Quoted",           "Expired / Declined",                 "Cancelled",       "System"],
            ["Bound",            "Documents generated",                "Issued",          "System"],
            ["Bound",            "Bind voided",                        "Cancelled",       "Admin"],
            ["Issued",           "Effective date reached",             "Active",          "System"],
            ["Issued",           "Pre-effective cancellation",         "Cancelled",       "Admin"],
            ["Active",           "Endorsement applied",                "Endorsed",        "Operations"],
            ["Endorsed",         "Endorsement takes effect",           "Active",          "System"],
            ["Active",           "60 days before expiry",              "Renewal Pending", "System"],
            ["Renewal Pending",  "Customer accepts renewal",           "Active (new)",    "Operations"],
            ["Renewal Pending",  "Renewal not accepted",               "Lapsed",          "System"],
            ["Active",           "Non-payment past grace period",      "Lapsed",          "System"],
            ["Lapsed",           "Reinstatement approved + payment",   "Reinstated",      "Admin"],
            ["Reinstated",       "Reinstatement effective",            "Active",          "System"],
            ["Active",           "Mid-term cancellation",              "Cancelled",       "Admin/Ops"],
            ["Active / Lapsed",  "End date reached",                   "Expired",         "System"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "4.2 Claim Lifecycle States")
    table(doc,
        ["From State", "Event / Trigger", "To State", "Actor"],
        [
            ["—",              "FNOL filed by lender",            "FNOL Submitted",  "Lender / Ops"],
            ["FNOL Submitted", "Assigned to claims handler",       "FNOL Processing", "System"],
            ["FNOL Processing","Claim formally registered",        "Filed",           "Claims Handler"],
            ["Filed",          "Investigation begins",             "Under Review",    "Claims Handler"],
            ["Under Review",   "Detailed fact-finding needed",     "Investigation",   "Claims Handler"],
            ["Under Review",   "Initial review complete",          "Evaluation",      "Claims Handler"],
            ["Investigation",  "Report ready",                     "Evaluation",      "Claims Handler"],
            ["Evaluation",     "Liability clear, amount agreed",   "Approved",        "Claims Handler"],
            ["Evaluation",     "Amount disputed",                  "Negotiation",     "Claims Handler"],
            ["Evaluation",     "Fraud / no valid claim",           "Rejected",        "Claims Handler"],
            ["Negotiation",    "Settlement terms agreed",          "Approved",        "Claims Handler"],
            ["Negotiation",    "Negotiation failed",               "Rejected",        "Claims Handler"],
            ["Approved",       "Payment disbursed",                "Settled",         "Claims Handler"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "4.3 Underwriting Decision States")
    table(doc,
        ["From State", "Event", "To State", "Condition"],
        [
            ["—",              "Policy submitted",               "Pending",                   "DDE complete"],
            ["Pending",        "Rules evaluated",                "Auto-Approved",              "All rules pass, LTV < threshold"],
            ["Pending",        "Referral rule triggered",        "Referred",                  "Risk score in referral band"],
            ["Pending",        "Hard rejection rule hit",        "Auto-Rejected",             "LTV > max or credit < min"],
            ["Referred",       "Assigned to senior UW",          "Pending Review",            "Within SLA (48h)"],
            ["Pending Review", "UW approves with riders",        "Approved with Conditions",  "Within authority limit"],
            ["Pending Review", "Exceeds UW authority",           "Escalated",                 "Coverage > maxCoverageAmount"],
            ["Pending Review", "UW declines",                    "Rejected",                  "Risk unacceptable"],
            ["Escalated",      "Senior authority approves",      "Approved with Conditions",  "Senior review complete"],
            ["Escalated",      "Senior authority declines",      "Rejected",                  "Risk unacceptable"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "4.4 Invoice Payment States")
    table(doc,
        ["From State", "Event", "To State"],
        [
            ["—",              "Invoice generated",               "Pending"],
            ["Pending",        "Full payment received",            "Paid"],
            ["Pending",        "Partial payment received",         "Partially Paid"],
            ["Pending",        "Due date passed, no payment",      "Overdue"],
            ["Partially Paid", "Remaining balance cleared",        "Paid"],
            ["Partially Paid", "Remaining balance overdue",        "Overdue"],
            ["Overdue",        "Late payment + fee received",      "Paid"],
            ["Overdue",        "Write-off approved",               "Void"],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5 — SEQUENCE DIAGRAMS
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "5. Sequence Diagrams")
    callout(doc,
        "Full Mermaid sequenceDiagram definitions are in APPLICATION_FLOW_DESIGN.md. "
        "This section provides step-by-step narrative descriptions of each flow.")

    flows = [
        ("5.1 Authentication Flow", [
            "User submits email and password on login page.",
            "Frontend sends POST /api/v1/auth/login to API Gateway.",
            "Rate limiter checks per-IP request count.",
            "Auth Service queries User DB to find user by email.",
            "bcrypt.compare() validates submitted password against stored hash.",
            "On success: JWT access token (15 min) and refresh token (7 days) generated.",
            "Refresh token stored in Redis keyed by userId.",
            "Tokens returned to frontend; user redirected to role-based dashboard.",
            "On failure: 401 Unauthorized returned with no detail disclosure.",
        ]),
        ("5.2 Token Refresh Flow", [
            "Frontend detects access token expiry (401 or timer).",
            "Sends POST /api/v1/auth/refresh with refresh token.",
            "Auth Service looks up refresh token in Redis.",
            "If not found / expired → 401, redirect to login.",
            "If valid → new access token issued, refresh token rotated.",
            "Updated tokens stored silently without user interruption.",
        ]),
        ("5.3 Policy Creation — QDE → DDE → Underwriting → Issuance", [
            "Phase 1 — QDE: Operations user completes quick data entry (customer, product, estimated coverage). Policy created in Draft status.",
            "Phase 2 — DDE: Detailed financial data entered (income, property value, LTV). Premium calculated using product rating factors. Status advances to Quoted. Customer notified.",
            "Phase 3 — Underwriting: UW service loads active rules in priority order. Each rule evaluated against applicant data. Risk score computed.",
            "  If Auto-Approve: policy moves to Bound. Operations notified.",
            "  If Refer: referral record created, SLA set to 48h, senior UW notified.",
            "  If Reject: policy cancelled, operations notified.",
            "Phase 4 — Issuance: Operations confirms issuance. Policy moves to Issued. Billing account and first invoice created. Policy Declaration document generated. Customer notified with policy number.",
        ]),
        ("5.4 Underwriting Referral Resolution", [
            "Referral appears in the UW queue dashboard.",
            "Underwriter opens referral — full risk context loaded (rules triggered, risk score, property data, credit score).",
            "Underwriter reviews and enters decision:",
            "  Accept: referral resolved, policy advanced to Bound, Operations notified.",
            "  Decline: policy cancelled, Operations notified.",
            "  Escalate: referral reassigned to senior authority, new notification sent.",
            "All actions recorded in audit log with timestamp and actor.",
        ]),
        ("5.5 FNOL Intake and Claims Registration", [
            "Lender / Operations user opens FNOL form.",
            "System validates the referenced policy is Active before allowing FNOL.",
            "FNOL record created (status=Submitted). Claims team notified.",
            "Claims handler processes FNOL: formal Claim registered with unique claim number.",
            "Fraud scoring engine triggered automatically.",
            "Initial reserve amount set based on estimated loss.",
            "FNOL status updated to 'Claim Created'. Lender notified with claim reference.",
        ]),
        ("5.6 Claim Adjudication and Settlement", [
            "Claims handler reviews full claim context (policy, FNOL, reserves, fraud score).",
            "Adjudication status progressed: Filed → Under Review → Investigation → Evaluation.",
            "Reserve adjusted after physical/document assessment.",
            "Handler enters approval decision and settlement amount.",
            "Settlement Offer document generated and sent to lender.",
            "On settlement confirmation: Claim status set to Settled, settlement date recorded.",
            "Billing service records claim payment in ledger. Settlement Letter generated. Audit log written.",
        ]),
        ("5.7 Billing — Invoice and Payment", [
            "Scheduled job identifies billing accounts with due dates within 7 days.",
            "Invoices generated automatically. Ledger debit entries created. Customers notified.",
            "Operations records payment received from lender (amount, method, reference).",
            "Payment record created. Invoice marked Paid. Billing account balance updated. Ledger credit entry added.",
            "Payment receipt notification sent to customer.",
            "Separate scheduled job detects overdue invoices → status updated → policy flagged as Delinquent.",
        ]),
        ("5.8 Policy Endorsement Flow", [
            "Operations user selects an Active policy and creates an endorsement request.",
            "Endorsement type selected (Coverage Change, Premium Adjustment, Beneficiary Change, Term Extension, Other).",
            "System checks actor's delegated authority limit against the requested coverage change.",
            "If within authority: endorsement auto-approved, policy updated (version incremented), billing adjusted, Endorsement Schedule document generated.",
            "If exceeds authority: endorsement held as Pending, senior approval notification sent.",
        ]),
        ("5.9 Policy Renewal Flow", [
            "Scheduled job triggers 60 days before each Active policy's expiry date.",
            "UW engine re-evaluates the policy using latest applicant data and claim history.",
            "New premium calculated based on updated risk score.",
            "Renewal record created (status=Quoted). Policy status set to Renewal Pending. Renewal Notice document generated. Customer and Operations notified.",
            "On acceptance: new policy record created (linkedto original via renewalOf). Original policy set to Expired. New billing account and invoice created.",
            "Customer notified with new policy number.",
        ]),
        ("5.10 Document Generation Flow", [
            "Operations user requests a specific document type (Certificate of Insurance, Endorsement Schedule, etc.).",
            "Document Service loads matching template definition from database.",
            "Entity data (policy / claim / invoice) fetched and merged into template.",
            "Rendered document uploaded to secure file store (S3 / Azure Blob).",
            "Document record created in database (type, version, filename, size).",
            "Audit log entry written. Download URL returned to user.",
        ]),
        ("5.11 Bulk Operations Flow", [
            "Admin selects multiple entities and triggers a bulk operation (Bulk Renewal, Bulk Cancel, Bulk Invoice, etc.).",
            "Bulk operation record created (status=Pending). 202 Accepted returned immediately with operation tracking ID.",
            "Background worker processes each entity asynchronously.",
            "Success and failure counts updated in real time.",
            "On completion: operation status set to Completed. Admin notified with summary (success count, failure count, per-entity results).",
            "Admin can query operation status at any time via GET /api/v1/bulk/operations/:id.",
        ]),
    ]

    for flow_title, flow_steps in flows:
        h2(doc, flow_title)
        for step in flow_steps:
            bullet(doc, step, level=(1 if step.startswith("  ") else 0))
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6 — DATA FLOW
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "6. Data Flow Diagrams")

    h2(doc, "6.1 End-to-End Policy Data Flow")
    stages = [
        ("Lender / Borrower", "Loan origination data input — borrower identity, property, loan amount"),
        ("QDE (Service Desk Intake)", "Customer identity, product selection, estimated coverage and loan amount"),
        ("DDE (Detailed Data Entry)", "Borrower financial profile, property details, LTV calculation, document uploads"),
        ("Underwriting Engine", "Rule evaluation (18+ active rules), risk score calculation, decision output"),
        ("Policy Issuance", "Policy number, Policy Declaration document, Certificate of Insurance, billing account, installment plan"),
        ("Billing", "Invoices generated per installment schedule, payments recorded, ledger maintained"),
        ("Servicing", "NPA tracking, delinquency monitoring, premium payment checks"),
        ("Claims Processing", "FNOL intake, claim registration, reserve setting, fraud scoring, adjudication, settlement payment"),
    ]
    for stage, desc in stages:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"{stage}:  "); r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(desc); r2.font.size = Pt(10)

    doc.add_paragraph()
    h2(doc, "6.2 Data Classification and Protection")
    table(doc,
        ["Classification", "Data Elements", "Protection Measures"],
        [
            ["Restricted",    "PAN, Aadhaar, Bank Account, Password Hash, Credit Score, Settlement Amounts",
             "AES-256 column encryption, RBAC-gated, never logged in plain text"],
            ["Confidential",  "Policy premium, Coverage amount, Claim amount, UW decision, Risk score",
             "RBAC-enforced access, masked in lower environments"],
            ["Internal",      "Policy status, Task assignments, Audit logs, Notifications, Reports",
             "Authentication required, role-filtered responses"],
            ["Public",        "Product descriptions, coverage limits, contact information",
             "No authentication required for general information"],
        ]
    )
    doc.add_paragraph()
    table(doc,
        ["Layer", "Protection Mechanism"],
        [
            ["Transport",  "TLS 1.3 for all client-server and service-service communication"],
            ["At Rest",    "AES-256 encryption for Restricted data fields at database column level"],
            ["Logging",    "PII fields masked / redacted before writing to log aggregator"],
            ["Audit",      "All state mutations recorded in immutable audit_logs table"],
            ["Backup",     "Encrypted database backups, tested restore procedures, 30-day retention"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "6.3 Integration Data Flow")
    table(doc,
        ["Integration", "Direction", "Protocol", "Data", "Resilience"],
        [
            ["Credit Bureau (CIBIL/Experian)", "Inbound",  "REST / SOAP", "Credit score, report",    "Circuit breaker, 24h Redis cache, 5s timeout"],
            ["Property Registry / Valuation",  "Inbound",  "REST",        "Property value, zone",    "Circuit breaker, retry x3, cached response"],
            ["SMS Gateway (Twilio)",           "Outbound", "REST",        "OTP, notifications",       "Retry queue, DLQ for failed sends"],
            ["Email Service (AWS SES)",        "Outbound", "SMTP / REST", "Invoices, policy docs",   "Retry x3, delivery tracking"],
            ["Regulatory Reporting API",       "Outbound", "REST / XML",  "Portfolio reports, filings","Idempotent submissions, ack tracking"],
            ["Document Store (S3 / Azure)",    "Outbound", "SDK",         "Generated documents, uploads","Pre-signed URLs, server-side encryption"],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7 — API REFERENCE
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "7. API Design Reference")

    h2(doc, "7.1 Conventions")
    bullets_api = [
        "Base URL: /api/v1",
        "GET    /resources                  → Paginated list with filters",
        "GET    /resources/:id              → Single resource detail",
        "POST   /resources                  → Create new resource",
        "PATCH  /resources/:id              → Partial update",
        "DELETE /resources/:id              → Soft delete (status='Deleted', record retained for audit)",
        "Sub-resources: POST /policies/:id/endorsements | /policies/:id/renewals | /claims/:id/reserves",
    ]
    for b in bullets_api:
        bullet(doc, b)

    doc.add_paragraph()
    h2(doc, "7.2 Standard Response Envelope")
    code_block(doc, """{
  "success": true,
  "data": { ... },
  "meta": { "page": 1, "limit": 20, "total": 245, "totalPages": 13 },
  "timestamp": "2026-03-07T10:30:00Z",
  "requestId": "req_abc123"
}

{
  "success": false,
  "error": {
    "code": "VALIDATION_ERROR",
    "message": "Request validation failed",
    "details": [{ "field": "coverageAmount", "message": "Must be between 100000 and 10000000" }]
  },
  "timestamp": "2026-03-07T10:30:00Z",
  "requestId": "req_abc123"
}""")

    doc.add_paragraph()
    h2(doc, "7.3 API Endpoint Catalogue")
    table(doc,
        ["Method", "Endpoint", "Auth", "Role", "Description"],
        [
            ["POST",  "/auth/login",                       "❌", "—",           "User login"],
            ["POST",  "/auth/refresh",                     "❌", "—",           "Refresh access token"],
            ["POST",  "/auth/logout",                      "✅", "All",         "Invalidate session"],
            ["GET",   "/policies",                         "✅", "All",         "List policies (paged, filtered)"],
            ["POST",  "/policies",                         "✅", "Admin, Ops",  "Create policy draft (QDE)"],
            ["GET",   "/policies/:id",                     "✅", "All",         "Get policy detail"],
            ["PATCH", "/policies/:id",                     "✅", "Admin, Ops",  "Update policy (DDE / status)"],
            ["POST",  "/policies/:id/issue",               "✅", "Admin, Ops",  "Issue policy"],
            ["POST",  "/policies/:id/cancel",              "✅", "Admin, Ops",  "Cancel policy"],
            ["POST",  "/policies/:id/reinstate",           "✅", "Admin, Ops",  "Reinstate lapsed policy"],
            ["POST",  "/policies/:id/endorsements",        "✅", "Admin, Ops",  "Create endorsement"],
            ["POST",  "/policies/:id/renewals",            "✅", "Admin, Ops",  "Initiate renewal"],
            ["POST",  "/underwriting/evaluate",            "✅", "Admin, UW",   "Run UW evaluation"],
            ["GET",   "/underwriting/referrals",           "✅", "Admin, UW",   "List referrals"],
            ["PATCH", "/underwriting/referrals/:id",       "✅", "Admin, UW",   "Resolve referral"],
            ["POST",  "/claims/fnol",                      "✅", "Admin, Ops, Claims", "File FNOL"],
            ["GET",   "/claims",                           "✅", "Admin, Ops, Claims", "List claims"],
            ["GET",   "/claims/:id",                       "✅", "All",         "Get claim detail"],
            ["PATCH", "/claims/:id",                       "✅", "Admin, Claims","Update claim"],
            ["POST",  "/claims/:id/reserves",              "✅", "Admin, Claims","Adjust reserve"],
            ["GET",   "/billing/accounts",                 "✅", "Admin, Ops",  "List billing accounts"],
            ["POST",  "/billing/payments",                 "✅", "Admin, Ops",  "Record payment"],
            ["GET",   "/billing/ledger/:accountId",        "✅", "Admin, Ops",  "Get ledger entries"],
            ["POST",  "/documents/generate",               "✅", "Admin, Ops",  "Generate document"],
            ["GET",   "/documents/:id",                    "✅", "All",         "Download document"],
            ["GET",   "/reports/dashboard",                "✅", "All",         "KPI dashboard data"],
            ["GET",   "/reports/loss-ratio",               "✅", "Admin, Ops",  "Loss ratio report"],
            ["POST",  "/bulk/operations",                  "✅", "Admin",       "Initiate bulk operation"],
            ["GET",   "/bulk/operations/:id",              "✅", "Admin",       "Track bulk operation status"],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8 — DEPLOYMENT
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, "8. Deployment Architecture")

    h2(doc, "8.1 Container Architecture (Kubernetes)")
    table(doc,
        ["Component", "Image", "Replicas", "Resources", "Notes"],
        [
            ["Frontend",          "pas-frontend:latest",  "2",   "0.5 CPU / 512MB",  "Next.js static export behind CDN"],
            ["Backend API",       "pas-backend:latest",   "3",   "1 CPU / 1GB",       "HPA: 3–10 pods based on CPU > 70%"],
            ["Background Workers","pas-workers:latest",   "2",   "0.5 CPU / 512MB",  "BullMQ jobs: renewal, billing, notifications"],
            ["PostgreSQL",        "postgres:15-alpine",   "1+1R","2 CPU / 4GB",       "Primary + read replica; PVC-backed"],
            ["Redis",             "redis:7-alpine",       "3",   "0.5 CPU / 1GB",    "Redis Cluster mode; persistence enabled"],
            ["NGINX Ingress",     "ingress-nginx",        "2",   "0.25 CPU / 256MB", "TLS termination, WAF rules, rate limit"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "8.2 CI/CD Pipeline Stages")
    table(doc,
        ["Stage", "Trigger", "Steps", "Gate"],
        [
            ["CI — Validation",   "Every push to PR",    "Lint, TypeScript check, Unit tests, Integration tests, SAST, Dependency audit, Docker build, Image scan", "All must pass"],
            ["Code Review",       "CI passes",           "Minimum 1 approval required; 2 required for security-sensitive paths", "Human approval"],
            ["CD — Staging",      "Merge to main",       "Helm deploy, DB migrations (Flyway), E2E tests (Playwright), DAST (OWASP ZAP), Performance test (k6)", "All must pass"],
            ["CD — Production",   "Manual gate",         "Blue/Green deploy, Smoke tests, Synthetic monitoring validation, Rollback standby", "Explicit approval"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "8.3 Observability Stack")
    table(doc,
        ["Pillar", "Tool", "What It Captures"],
        [
            ["Metrics",    "Prometheus + Grafana",          "API request rate, error %, latency percentiles, DB pool utilisation, pod CPU/memory"],
            ["Logs",       "ELK Stack / OpenSearch",        "Structured JSON logs from all services; PII fields masked; 90-day retention"],
            ["Traces",     "Jaeger (OpenTelemetry)",         "Distributed request traces across frontend → backend → database → external APIs"],
            ["Alerting",   "PagerDuty / OpsGenie",          "P1: 15min, P2: 30min response SLAs; on-call rotation; escalation paths"],
            ["Dashboards", "Grafana Business Dashboard",    "Loss ratio, combined ratio, claims frequency, billing aging, policy count by status"],
            ["Synthetic",  "Checkly / Datadog Synthetics",  "Critical user journeys (login, policy create, claim file) tested every 5 minutes"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "8.4 Disaster Recovery")
    table(doc,
        ["Parameter", "Target"],
        [
            ["Recovery Time Objective (RTO)",    "< 4 hours for P1 production outage"],
            ["Recovery Point Objective (RPO)",   "< 1 hour (PostgreSQL WAL streaming to standby)"],
            ["Database Backup Frequency",         "Daily full + continuous WAL archiving"],
            ["Backup Retention",                  "30 days on-site, 12 months cold storage"],
            ["Multi-Region Failover",             "Secondary region active standby; DNS failover < 5 min"],
            ["Rollback Capability",               "Application: Kubernetes rollback < 5 min; DB: migration rollback scripts"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Document maintained by the Enterprise Architecture team.  "
                   "Version 1.0 | March 2026 | Review: Per major release or architectural change.")
    r.italic = True; r.font.size = Pt(9)

    out = "/home/user/PAS_POC/docs/APPLICATION_FLOW_DESIGN.docx"
    doc.save(out)
    print(f"Saved: {out}")

build()
