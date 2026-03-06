"""
Script to generate AI_ASSISTED_DEVELOPMENT_GUIDELINES.docx
using the python-docx library.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set background colour for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(table, border_color="CCCCCC"):
    """Apply thin borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color="1B3A6B"):
    """Add a styled heading."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return para


def add_para(doc, text, bold=False, italic=False, size=10, indent=False, color=None):
    """Add a normal paragraph."""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = para.add_run(text)
    run.font.size = Pt(10)
    return para


def add_table(doc, headers, rows, header_bg="1B3A6B", header_fg="FFFFFF", alt_bg="EBF0FA"):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(header_fg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_bg if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)

    set_cell_border(table)
    return table


def add_callout(doc, text, bg="FFF3CD", border="F0AD4E"):
    """Add a highlighted callout / note box."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    set_cell_border(table, border)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default body font ─────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ======================================================================
    # COVER PAGE
    # ======================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AI-Assisted Development Guidelines\nfor Enterprise Applications")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string("1B3A6B")

    doc.add_paragraph()
    meta_lines = [
        ("Document Classification:", "Internal — Technical Leadership & Management"),
        ("Version:", "1.0"),
        ("Date:", "March 2026"),
        ("Prepared By:", "Enterprise Architecture & AI Engineering"),
        ("Audience:", "Technical Leadership, Enterprise Architects, Engineering Management, Compliance & Security Teams"),
    ]
    for label, value in meta_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(f"{label} ")
        r.bold = True
        r.font.size = Pt(10)
        r2 = para.add_run(value)
        r2.font.size = Pt(10)

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS (static)
    # ======================================================================
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Overview of AI-Assisted Development",
        "2. Code Quality Standards",
        "3. Enterprise-Grade Capabilities of Claude Code Generated Applications",
        "4. Use of Loveable for UI Prototyping",
        "5. Risk Assessment of AI-Generated Code",
        "6. Production Support Strategy",
        "7. Bug Handling Strategy After Production Deployment",
        "8. Governance Model for AI-Assisted Development",
        "Appendix A: Glossary",
        "Appendix B: Recommended Tooling Reference",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    # ======================================================================
    # SECTION 1 — Overview
    # ======================================================================
    add_heading(doc, "1. Overview of AI-Assisted Development", level=1)

    add_heading(doc, "1.1 What Is AI-Assisted Development?", level=2)
    add_para(doc,
        "AI-assisted development refers to the integration of large language model (LLM)-powered tools into "
        "the software development lifecycle (SDLC) to augment developer productivity, accelerate delivery, "
        "and improve consistency. Tools such as Claude Code (developed by Anthropic) operate as intelligent "
        "engineering collaborators capable of reading, writing, refactoring, debugging, and reviewing code "
        "across the full application stack.")
    doc.add_paragraph()
    add_para(doc,
        "Unlike traditional code generation tools or simple autocomplete engines, modern AI coding assistants "
        "like Claude Code reason contextually over entire codebases, understand architectural constraints, "
        "propose solutions consistent with existing patterns, and explain trade-offs at a level appropriate "
        "for experienced engineers.")

    add_heading(doc, "1.2 Role of Claude Code in the Development Lifecycle", level=2)
    add_table(doc,
        ["SDLC Phase", "Claude Code Contribution"],
        [
            ["Requirements Analysis",  "Translates business requirements into technical specifications and data models"],
            ["Architecture Design",    "Proposes and evaluates architectural patterns (microservices, event-driven, layered)"],
            ["Implementation",         "Generates production-quality code with tests, error handling, and documentation"],
            ["Code Review",            "Identifies anti-patterns, security vulnerabilities, and deviations from standards"],
            ["Testing",                "Writes unit, integration, and end-to-end test suites"],
            ["Debugging",              "Diagnoses runtime errors and traces root causes across call stacks"],
            ["Documentation",          "Produces inline comments, API documentation, and architectural decision records"],
            ["Refactoring",            "Modernises legacy code while preserving functional behaviour"],
            ["DevOps & CI/CD",         "Generates pipeline configurations, Dockerfiles, and infrastructure-as-code"],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "1.3 Strategic Value for Enterprise Development", level=2)
    bullets_s1 = [
        "Velocity: AI assistance compresses development timelines significantly without sacrificing quality when governed correctly.",
        "Consistency: AI tools apply the same coding standards uniformly across all modules, reducing stylistic drift between developers.",
        "Knowledge Amplification: Junior developers gain access to senior-level patterns and practices embedded in AI recommendations.",
        "Risk Reduction: Proactive identification of security vulnerabilities, anti-patterns, and missing error handling during development rather than post-deployment.",
        "Cost Efficiency: Reduction in rework cycles, manual code reviews, and late-stage defect remediation.",
    ]
    for b in bullets_s1:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_callout(doc,
        "Management Takeaway: AI-assisted development is not a replacement for software engineers. It is a force multiplier. "
        "Human engineers remain responsible for architecture decisions, business logic validation, security governance, and code "
        "approval. Claude Code accelerates the mechanical aspects of software construction while humans retain full ownership of "
        "quality, correctness, and compliance.")
    doc.add_page_break()

    # ======================================================================
    # SECTION 2 — Code Quality
    # ======================================================================
    add_heading(doc, "2. Code Quality Standards", level=1)
    add_para(doc,
        "Code quality is not binary — it exists on a spectrum. In enterprise software, particularly systems "
        "that handle sensitive production data, understanding and enforcing code quality is a fundamental risk "
        "management obligation. Below are three classifications of code quality with their characteristics, "
        "risks, and relevance to AI-assisted development.")

    for category, color, sections in [
        ("2.1 Good Code",  "1B7A2F", [
            ("Characteristics", [
                "Readable and Self-Documenting: Variable names, function names, and class names clearly communicate intent.",
                "Single Responsibility: Each function, class, and module has one clearly defined purpose.",
                "Testable: Logic is decoupled from infrastructure concerns, enabling straightforward unit testing.",
                "Secure by Design: Input validation, output encoding, parameterised queries, and least privilege applied consistently.",
                "Maintainable: Changes can be made in one place without causing cascading failures.",
                "Consistent: Follows team-agreed coding standards, linting rules, and architectural conventions.",
                "Error Handling: Exceptions are handled gracefully, logged with sufficient context without exposing sensitive data.",
                "Performant: Algorithms and data structures are chosen appropriately; queries are optimised.",
            ]),
            ("Risks", [
                "Requires time and discipline investment — shortcuts under pressure gradually degrade quality.",
                "Good code today can become technical debt if surrounding architecture evolves without the module being updated.",
            ]),
            ("Example Scenario", [
                "A policy management API endpoint that validates all fields against a schema, enforces JWT authentication, "
                "performs parameterised queries, logs all state changes to an audit trail, and returns standardised error "
                "responses. Every function is under 30 lines with 90%+ unit test coverage.",
            ]),
            ("How Claude Code Enforces This", [
                "Generates code adhering to SOLID principles by default when sufficient context is provided.",
                "Automatically includes input validation using established libraries (Zod, Joi, class-validator).",
                "Proposes and writes unit tests as part of the implementation task.",
                "Flags violations of best practices when reviewing code and suggests refactored alternatives.",
            ]),
        ]),
        ("2.2 Bad Code", "C0392B", [
            ("Characteristics", [
                "Unclear Intent: Variable names such as x, temp, data2 or functions named doStuff() with no semantic meaning.",
                "Long Methods / God Classes: Single functions spanning hundreds of lines performing multiple unrelated operations.",
                "No Error Handling: Operations that can fail executed without try/catch blocks or fallback logic.",
                "Hardcoded Values: Credentials, URLs, and configuration parameters embedded directly in source code.",
                "Code Duplication: Same logic copy-pasted across multiple locations, causing inconsistent bug fixes.",
                "No Tests: Modules shipped without any automated test coverage.",
                "Tight Coupling: Business logic, database access, and presentation logic intermingled with no separation.",
            ]),
            ("Risks", [
                "Security Vulnerabilities: Hardcoded credentials and missing validation are direct pathways to SQL injection, XSS, and credential leakage.",
                "Maintenance Burden: Every change requires deciphering opaque, duplicated logic, dramatically increasing cost.",
                "Regression Risk: Without tests, breaking changes are detected only in production.",
                "Knowledge Dependency: When the original developer leaves, bad code becomes an organisational liability.",
            ]),
            ("Example Scenario", [
                "A claims processing function that connects to a database using a hardcoded password, executes dynamically "
                "constructed SQL using unvalidated user input (SQL injection risk), returns raw database error messages to "
                "the client, has no logging, and is called by 12 different parts of the application with no abstraction.",
            ]),
            ("How Claude Code Helps", [
                "Trained to avoid generating code with these anti-patterns for production use cases.",
                "When reviewing existing bad code, identifies anti-patterns and proposes refactored alternatives.",
                "Must be provided appropriate context — without quality guidelines, optimises for brevity over production readiness.",
            ]),
        ]),
        ("2.3 Ugly Code", "D68910", [
            ("Characteristics", [
                "Inconsistent Formatting: Mixed indentation styles and naming conventions within the same codebase.",
                "Commented-Out Dead Code: Blocks left in place creating noise and confusion.",
                "Over-Engineered Abstractions: Unnecessary inheritance hierarchies or complex patterns for trivially simple problems.",
                "Under-Engineered Solutions: Complex logic handled with deeply nested if/else chains.",
                "Inconsistent Error Handling: Some paths handle errors, others do not.",
                "Technical Debt Comments: TODO, FIXME, and HACK comments indicating known issues never addressed.",
            ]),
            ("Risks", [
                "Maintainability Decay: Gateway to bad code as complexity accumulates without refactoring.",
                "Onboarding Friction: New developers take longer to become productive.",
                "Bug Concealment: Inconsistent formatting makes bugs harder to spot during code review.",
                "Velocity Reduction: Developers spend increasing time understanding existing code rather than writing new features.",
            ]),
            ("Example Scenario", [
                "A billing module that works correctly but uses four different naming conventions, contains 200-line functions "
                "with 6 levels of nesting, has 40 lines of commented-out experimental code, and alternates between two "
                "different error-handling approaches. A new developer takes three days to understand it well enough to modify safely.",
            ]),
            ("How Claude Code Helps", [
                "Applies consistent formatting and naming conventions aligned with project standards when generating code.",
                "When reviewing ugly code, identifies inconsistencies and proposes normalised naming and refactoring strategies.",
                "Can be used as a dedicated refactoring tool in a sprint to systematically clean up technical debt.",
            ]),
        ]),
    ]:
        add_heading(doc, category, level=2)
        for sub_title, items in sections:
            add_para(doc, sub_title, bold=True, size=10)
            for item in items:
                add_bullet(doc, item)
        doc.add_paragraph()

    add_heading(doc, "2.4 Summary: Code Quality Comparison", level=2)
    add_table(doc,
        ["Dimension", "Good Code", "Bad Code", "Ugly Code"],
        [
            ["Readability",        "High",         "Low",      "Medium"],
            ["Testability",        "High",         "Low",      "Medium"],
            ["Security",           "Built-in",     "Absent",   "Inconsistent"],
            ["Maintainability",    "High",         "Low",      "Declining"],
            ["Onboarding Ease",    "Easy",         "Difficult","Difficult"],
            ["Production Risk",    "Low",          "High",     "Medium-High"],
            ["AI Remediation",     "Maintain & extend","Refactor immediately","Systematic clean-up"],
        ]
    )
    doc.add_page_break()

    # ======================================================================
    # SECTION 3 — Enterprise Capabilities
    # ======================================================================
    add_heading(doc, "3. Enterprise-Grade Capabilities of Claude Code Generated Applications", level=1)
    add_para(doc,
        "Enterprise software must meet a higher standard than prototypes or internal tools. It must be secure, "
        "scalable, compliant, and operable at production scale with sensitive data.")

    # 3.1 Security
    add_heading(doc, "3.1 Security", level=2)

    add_heading(doc, "3.1.1 Secure Coding Practices", level=3)
    for b in [
        "Input Validation: All external data validated against strict schemas before processing (Zod, Pydantic).",
        "Output Encoding: Data rendered to UI or returned in API responses is properly encoded to prevent XSS.",
        "Parameterised Queries: Database operations use prepared statements or ORM-level parameterisation, eliminating SQL injection.",
        "Secret Management: Credentials and API keys externalised to environment variables or vault-based secret stores.",
        "Principle of Least Privilege: Application service accounts granted only minimum required permissions.",
        "Dependency Security: Dependencies pinned and scanned using npm audit, pip-audit, or Snyk.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.1.2 OWASP Top 10 Coverage", level=3)
    add_table(doc,
        ["OWASP Risk", "Mitigation Applied"],
        [
            ["A01 Broken Access Control",     "RBAC middleware, JWT validation, resource-level ownership checks"],
            ["A02 Cryptographic Failures",    "TLS enforced, bcrypt/Argon2 for passwords, AES-256 for data at rest"],
            ["A03 Injection",                 "Parameterised queries, schema validation, ORM usage"],
            ["A04 Insecure Design",           "Security requirements modelled during design, threat modelling applied"],
            ["A05 Security Misconfiguration", "Secure defaults, no debug modes in production, security headers via Helmet.js"],
            ["A06 Vulnerable Components",     "Dependency scanning in CI/CD, automated CVE alerts"],
            ["A07 Auth & Session Failures",   "JWT with short expiry, refresh token rotation, brute force protection"],
            ["A08 Software Integrity Failures","Signed commits, verified package checksums, supply chain scanning"],
            ["A09 Logging & Monitoring Failures","Centralised structured logging, audit trails, real-time alerting"],
            ["A10 SSRF",                      "Allowlisted outbound URLs, request validation for server-side HTTP calls"],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "3.1.3 Authentication and Authorisation", level=3)
    for b in [
        "Authentication: JWT-based stateless authentication with configurable expiry. MFA integration for privileged accounts.",
        "Authorisation: Granular RBAC enforced at the API middleware layer. All decisions are server-side.",
        "Session Management: Secure, HttpOnly, SameSite cookies. CSRF tokens. Automatic session expiry.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.1.4 Data Protection", level=3)
    for b in [
        "Encryption in Transit: All communications use TLS 1.2 or higher. Internal services use mTLS where feasible.",
        "Encryption at Rest: Sensitive fields encrypted at database level using AES-256.",
        "Data Masking: PII masked in logs, API responses to non-privileged roles, and test environments.",
        "Retention Policies: Automated purge jobs aligned to regulatory requirements.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.1.5 Secure API Design", level=3)
    for b in [
        "API Versioning: All endpoints versioned (/api/v1/, /api/v2/) to allow breaking changes without disrupting integrations.",
        "Rate Limiting: Per-user and per-IP rate limiting enforced at API gateway level.",
        "Input Size Limits: Request payload size limits enforced to prevent memory exhaustion.",
        "CORS Policy: Strict policy with explicit allowlist of trusted origins.",
        "API Gateway Security: Authentication, rate limiting, logging, and threat detection before reaching application services.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()

    # 3.2 Scalability
    add_heading(doc, "3.2 Scalability", level=2)

    add_heading(doc, "3.2.1 Microservices Readiness", level=3)
    for b in [
        "Domain-Driven Design (DDD): Business domains modelled as bounded contexts with clear interface contracts.",
        "Loose Coupling: Services communicate through well-defined interfaces. No direct cross-domain database calls.",
        "Independent Deployability: Each service can be built, tested, deployed, and scaled independently.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.2.2 Horizontal Scalability", level=3)
    for b in [
        "Stateless Services: No user session state on instances. All state stored externally (database, cache).",
        "Load Balancing: Traffic distributed across multiple instances with health checks.",
        "Database Scalability: Read replicas, sharding strategies, connection pooling.",
        "Caching Layer: Redis for frequently accessed, low-volatility data.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.2.3 Cloud-Native Architecture", level=3)
    for b in [
        "Infrastructure as Code (IaC): All infrastructure defined in code (Terraform, CloudFormation) and version-controlled.",
        "Managed Services: Database, caching, queuing, and secret management via managed cloud services.",
        "Auto-Scaling: HPA configured for automatic scale-out under load.",
        "Multi-Region: Deployable across multiple geographic regions for disaster recovery.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.2.4 Containerisation Readiness", level=3)
    for b in [
        "Docker: Application services packaged as Docker images with security-hardened base images.",
        "Kubernetes: Deployment manifests, services, ingress, and secrets defined as YAML, ready for EKS/AKS/GKE.",
        "Twelve-Factor Compliance: Configuration separated from code, state externalised, environment-specific config via env vars.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.2.5 Performance Considerations", level=3)
    for b in [
        "Database Query Optimisation: Queries reviewed and optimised. Indexes created based on query patterns.",
        "Asynchronous Processing: Long-running operations offloaded to background job queues.",
        "Pagination: All list endpoints return paginated results with configurable page sizes.",
        "Response Compression: HTTP response compression (gzip/brotli) enabled to reduce bandwidth.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()

    # 3.3 Compliance
    add_heading(doc, "3.3 Compliance", level=2)

    add_heading(doc, "3.3.1 Handling Sensitive Production Data", level=3)
    for b in [
        "Data Classification: All data classified by sensitivity (Public, Internal, Confidential, Restricted) with corresponding controls.",
        "PII Handling: PII protected through access controls, encryption, and audit logging. Never written to logs in plain text.",
        "Data Minimisation: Only data required for a specific function is collected, processed, and stored.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.3.2 Logging and Auditing", level=3)
    for b in [
        "Application Logs: Structured JSON logging captures event type, timestamp, user identity, resource, and outcome.",
        "Audit Trail: Every state-changing operation on sensitive entities generates an immutable audit log entry.",
        "Access Logs: All API access logged including user identity, endpoint, method, response code, and response time.",
        "Log Retention: Retention periods configured per regulatory requirements (e.g., 7 years for financial records).",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "3.3.3 Regulatory Frameworks Supported", level=3)
    add_table(doc,
        ["Framework", "Key Obligations"],
        [
            ["GDPR",     "Data subject rights, consent management, privacy by design, breach notification"],
            ["POPIA",    "South African personal data protection, information officer accountability"],
            ["PCI-DSS",  "Secure handling of payment card data, encrypted storage and transmission"],
            ["SOC 2",    "Security, availability, confidentiality, processing integrity, privacy controls"],
            ["ISO 27001","Information security management system (ISMS) alignment"],
        ]
    )
    doc.add_page_break()

    # ======================================================================
    # SECTION 4 — Loveable Prototyping
    # ======================================================================
    add_heading(doc, "4. Use of Loveable for UI Prototyping", level=1)

    add_heading(doc, "4.1 What Are UI Prototyping Tools?", level=2)
    add_para(doc,
        "UI prototyping tools such as Loveable (formerly GPT Engineer), v0 by Vercel, and Bolt are AI-powered "
        "rapid application development environments designed to transform natural language descriptions into "
        "functional user interface code at exceptional speed. These tools generate React, HTML/CSS, and sometimes "
        "basic backend scaffolding in minutes, enabling teams to rapidly visualise product ideas and gather "
        "stakeholder feedback before investing in production engineering effort.")

    add_heading(doc, "4.2 What Prototyping Tools Are Used For", level=2)
    for b in [
        "Stakeholder Demonstrations: Quickly build a clickable, visual representation of a product concept to facilitate business discussions.",
        "User Experience Validation: Test navigation flows, layout decisions, and feature interactions with end users before committing to technical implementation.",
        "Requirements Clarity: Use a tangible prototype to surface missing requirements, ambiguities, and design conflicts early.",
        "Design Handoff: Provide design references to engineering teams as a starting point for UI development.",
        "Investor and Client Presentations: Produce polished-looking demonstrations of proposed products.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "4.3 Prototype Code vs. Production-Grade Code", level=2)
    add_table(doc,
        ["Dimension", "Prototype Code (Loveable)", "Production-Grade Code"],
        [
            ["Purpose",          "Demonstrate concept, gather feedback",   "Serve real users with real data"],
            ["Security",         "None",                                    "Comprehensive — OWASP-compliant"],
            ["Data",             "Mock / hardcoded",                        "Live, encrypted, governed"],
            ["Error Handling",   "Minimal or absent",                       "Comprehensive, with logging"],
            ["Testing",          "None",                                    "Unit, integration, E2E test suites"],
            ["Scalability",      "Single-user, no concurrency",             "Horizontally scalable, load-tested"],
            ["Architecture",     "Monolithic, tightly coupled",             "Layered, microservices-ready"],
            ["Auditability",     "None",                                    "Full audit trail"],
            ["DevOps",           "None",                                    "Full CI/CD pipeline"],
            ["Compliance",       "None",                                    "Regulatory-aligned"],
            ["Maintainability",  "Low — built for speed",                   "High — built for longevity"],
        ]
    )
    doc.add_paragraph()
    add_callout(doc,
        "Critical Guidance for Management: Prototype code from Loveable should never be promoted directly to production. "
        "It is a design artefact, not a deployable product. Using prototype code with production data represents an "
        "unacceptable security and compliance risk.")

    add_heading(doc, "4.4 Can the Prototype Be Converted to Enterprise-Grade?", level=2)
    add_para(doc,
        "Yes, with a structured transformation programme. The prototype serves as a valuable reference for visual design "
        "and user experience decisions, feature scope, and navigation flow. However, the underlying code must be treated "
        "as a blueprint, not a foundation. The recommended approach is to use the prototype as a functional specification "
        "and rebuild the application on an enterprise-grade technical foundation.")

    add_heading(doc, "4.5 Transformation Roadmap", level=2)
    steps = [
        ("Step 1: Architecture Redesign", [
            "Define target architecture: layered or microservices based on scale requirements.",
            "Define domain boundaries: Policy, Claims, Billing, Underwriting, Customer Management.",
            "Design the data model with normalised, indexed relational or document database schema.",
            "Design RESTful API contracts with versioning and OpenAPI/Swagger documentation.",
            "Define cloud infrastructure using Infrastructure as Code.",
        ]),
        ("Step 2: Security Hardening", [
            "Implement JWT-based authentication or integrate with enterprise SSO (SAML, OIDC, Okta, Azure AD).",
            "Design and implement granular RBAC model aligned with business roles.",
            "Apply schema validation for all API endpoint inputs.",
            "Move all credentials to vault-based secret management.",
            "Apply HTTP security headers and conduct penetration testing before production launch.",
        ]),
        ("Step 3: Backend Integration", [
            "Replace mock data with real persistence using parameterised queries or ORM.",
            "Implement business logic layer in a proper service layer on the backend.",
            "Integrate external systems through properly abstracted integration adapters.",
            "Implement asynchronous workflows for notifications, reports, and batch operations.",
        ]),
        ("Step 4: Code Refactoring", [
            "Decompose monolithic components into focused, reusable components.",
            "Enforce linting rules (ESLint, Prettier) across the entire codebase.",
            "Replace all hardcoded values with environment-driven configuration.",
            "Eliminate dead code, unused imports, and unreachable branches.",
            "Apply SOLID principles throughout the service layer.",
        ]),
        ("Step 5: Performance Optimisation", [
            "Frontend: Code splitting, lazy loading, image optimisation, bundle analysis.",
            "Backend: Query analysis, index creation, connection pooling, response caching.",
            "Load testing using k6, Locust, or JMeter to identify and resolve bottlenecks.",
            "CDN integration for static assets.",
        ]),
        ("Step 6: Testing Framework", [
            "Unit tests: Jest/Vitest targeting 80%+ code coverage.",
            "Integration tests: Service, repository, and external API interaction tests.",
            "End-to-end tests: Complete user journeys using Playwright or Cypress.",
            "Security tests: Automated OWASP ZAP scanning integrated into the test suite.",
            "Performance tests: Load and stress tests as part of pre-release validation.",
        ]),
        ("Step 7: DevOps Pipeline", [
            "Git branching strategy with branch protection and mandatory code reviews.",
            "CI pipeline: lint, unit tests, integration tests, security scan, build on every PR.",
            "CD pipeline: automated deployment to staging on merge; production requires approval gate.",
            "Containerisation: Docker images deployed via Kubernetes.",
            "Monitoring and APM integrated from day one (Datadog, New Relic, CloudWatch).",
            "All cloud resources defined in Terraform, version-controlled, and deployed through the pipeline.",
        ]),
    ]
    for step_title, step_bullets in steps:
        add_para(doc, step_title, bold=True, size=10)
        for b in step_bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    doc.add_page_break()

    # ======================================================================
    # SECTION 5 — Risk Assessment
    # ======================================================================
    add_heading(doc, "5. Risk Assessment of AI-Generated Code", level=1)

    add_heading(doc, "5.1 Identified Risks", level=2)
    risks = [
        ("Code Quality Inconsistency",
         "AI tools generate code that varies in quality depending on instruction specificity. Without consistent "
         "prompting standards, generated code may vary in structure, naming conventions, and architectural pattern adherence.",
         "Maintenance complexity, inconsistent behaviour, increased debugging time."),
        ("Dependency Risks",
         "AI tools may recommend dependencies that are outdated, unmaintained, or have known vulnerabilities.",
         "Security vulnerabilities, licensing compliance risks, build failures from deprecated packages."),
        ("Security Vulnerabilities",
         "Without appropriate context, AI tools may omit security controls such as input validation, authentication checks, "
         "or proper error handling.",
         "Potential exposure of sensitive data, authentication bypass, injection vulnerabilities."),
        ("Lack of Architecture Discipline",
         "AI-generated code may not naturally respect defined architectural boundaries, introducing direct database calls "
         "in the wrong layer or tightly coupling modules that should be independent.",
         "Technical debt accumulation, reduced testability, difficulty scaling or refactoring."),
        ("Hallucination and Incorrect Logic",
         "LLMs can generate syntactically correct code that contains subtle logical errors or incorrect business rule "
         "implementations.",
         "Incorrect data processing, data integrity issues, potential regulatory violations."),
    ]
    for risk_title, risk_desc, risk_impact in risks:
        add_para(doc, f"5.1.x  {risk_title}", bold=True, size=10)
        add_para(doc, f"Risk: {risk_desc}")
        add_para(doc, f"Impact: {risk_impact}", italic=True)
        doc.add_paragraph()

    add_heading(doc, "5.2 Risk Mitigation Strategy", level=2)
    mitigations = [
        ("5.2.1  Mandatory Human Code Review", [
            "Every AI-generated code change must be reviewed and approved by a qualified human engineer before merging.",
            "Reviewers assess: business logic correctness, architectural compliance, security posture, test coverage, and documentation.",
            "Two-person review required for authentication code, payment processing, encryption routines, and CI/CD configuration.",
        ]),
        ("5.2.2  Static Analysis (SAST)", [
            "SonarQube / SonarCloud: Code quality gates enforce minimum standards. PRs failing quality gates are blocked.",
            "ESLint / TypeScript Strict Mode: Type safety and linting rules catch common errors at editor and CI pipeline level.",
            "Semgrep: Security-focused analysis identifying security anti-patterns specific to technologies used.",
        ]),
        ("5.2.3  Dependency Scanning", [
            "npm audit / pip-audit / Dependabot: Automated scanning for CVEs. Critical vulnerabilities block the build.",
            "Software Bill of Materials (SBOM): Generated for each release, providing a complete inventory of software components.",
            "License compliance scanning to identify conflicts with commercial software distribution.",
        ]),
        ("5.2.4  Dynamic Security Scanning (DAST)", [
            "OWASP ZAP / Burp Suite: Automated dynamic scanning against staging environment on each release.",
            "Identifies runtime security vulnerabilities that static analysis cannot detect.",
        ]),
        ("5.2.5  Architecture Governance", [
            "Architecture Decision Records (ADRs): All significant decisions documented and stored in the repository.",
            "Architecture Review Board: Changes affecting system architecture require team sign-off.",
            "Coding Standards Document: Shared with AI assistant as context to ensure generated code aligns with conventions.",
            "Automated linting and formatting enforced in pre-commit hooks and CI pipelines.",
        ]),
    ]
    for mit_title, mit_bullets in mitigations:
        add_para(doc, mit_title, bold=True, size=10)
        for b in mit_bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    doc.add_page_break()

    # ======================================================================
    # SECTION 6 — Production Support Strategy
    # ======================================================================
    add_heading(doc, "6. Production Support Strategy (When AI Tools Are Not Used)", level=1)
    add_para(doc,
        "If AI coding assistants are not available or not used during future support and maintenance phases, "
        "the following strategy ensures continued effective operation of the production system.")

    support_sections = [
        ("6.1 Traditional Developer-Led Debugging", [
            "Dedicated production support engineers assigned on a rotational basis.",
            "Support engineers trained on application architecture and common failure modes before taking production responsibility.",
            "Debugging follows structured escalation: L1 (monitoring and triage), L2 (investigation and workaround), L3 (root cause and fix).",
            "All debugging sessions time-boxed to prevent prolonged production incidents.",
        ]),
        ("6.2 Source Code Documentation", [
            "Module-level documentation: Each module includes a header comment explaining purpose and key dependencies.",
            "API documentation: All REST endpoints documented using OpenAPI/Swagger specifications.",
            "Architecture documentation: High-level diagrams and sequence diagrams for critical workflows in docs/ directory.",
            "Runbook library: Step-by-step operational runbooks for common support scenarios.",
            "Database schema documentation: Entity-relationship diagrams and data dictionary maintained with each schema change.",
        ]),
        ("6.3 Knowledge Transfer", [
            "Structured handover sessions: AI-assisted developers walk through critical modules and known issue patterns.",
            "Pair programming: New support engineers pair with experienced developers before taking independent responsibility.",
            "Video walkthroughs: Recorded technical walkthroughs of key system components stored in the knowledge base.",
            "Architecture Q&A sessions: Open sessions where support engineers can ask detailed technical questions.",
        ]),
        ("6.4 Code Ownership Model", [
            "Module ownership: Each module has an assigned primary owner and secondary owner (backup coverage).",
            "CODEOWNERS file: Automatically assigns appropriate reviewers to pull requests affecting each module.",
            "Ownership documentation maintained and updated when team membership changes.",
        ]),
        ("6.5 Manual Debugging and Monitoring Tools", [
            "Log analysis: Centralised logging platform provides search, filtering, and visualisation for investigation.",
            "Application Performance Monitoring (APM): Distributed tracing and performance metrics without code-level analysis.",
            "Database query analysis: Slow query logs, EXPLAIN plans, and database-level monitoring for data layer issues.",
            "Network analysis: API gateway and network flow logs for connectivity and integration issues.",
        ]),
    ]
    for sec_title, sec_bullets in support_sections:
        add_heading(doc, sec_title, level=2)
        for b in sec_bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    doc.add_page_break()

    # ======================================================================
    # SECTION 7 — Bug Handling
    # ======================================================================
    add_heading(doc, "7. Bug Handling Strategy After Production Deployment", level=1)

    add_heading(doc, "7.1 Production Monitoring", level=2)
    for b in [
        "Synthetic monitoring: Automated health check scripts execute representative user journeys on a schedule.",
        "Real User Monitoring (RUM): Frontend performance and error rates monitored from actual user perspective.",
        "Infrastructure monitoring: CPU, memory, disk, and network metrics with automated threshold alerts.",
        "Application metrics: Custom business-level metrics tracking throughput anomalies (policies processed/hour, claims API error rate).",
        "Error rate alerting: Automatic alerts when error rates on any endpoint exceed defined threshold (e.g., 1% 5xx over 5 minutes).",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "7.2 Incident Classification", level=2)
    add_table(doc,
        ["Severity", "Definition", "Response Time", "Resolution Target"],
        [
            ["P1 — Critical", "Production outage or data breach. All users affected.",     "15 minutes", "4 hours"],
            ["P2 — High",     "Major feature unavailable. Significant user impact.",        "30 minutes", "8 hours"],
            ["P3 — Medium",   "Degraded functionality. Workaround available.",              "2 hours",    "24 hours"],
            ["P4 — Low",      "Minor issue. Minimal user impact.",                          "Next business day", "Next sprint"],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "7.3 Logging and Observability", level=2)
    for b in [
        "Logs: Structured JSON logs from every component, centralised and searchable. Correlation IDs trace requests across services.",
        "Metrics: Request rate, error rate, latency percentiles, and database connection pool utilisation in dashboards and alerts.",
        "Traces: Distributed tracing (OpenTelemetry) captures the complete execution path across all services.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "7.4 Root Cause Analysis", level=2)
    for b in [
        "Formal RCA for all P1/P2 incidents following '5 Whys' or Ishikawa methodology.",
        "Timeline reconstruction using logs and metrics to identify exact sequence of events.",
        "Contributing factor identification covering both technical and process factors.",
        "Specific, measurable corrective actions assigned to owners and tracked to completion.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "7.5 Hotfix and Patch Deployment", level=2)
    for b in [
        "Hotfix branch process: Critical fixes developed on dedicated hotfix/ branch from current production tag.",
        "Expedited review: Minimum one senior engineer sign-off. Security review mandatory for security-sensitive changes.",
        "Fast-track pipeline: Same validation checks as normal deployments, prioritised in queue.",
        "CAB fast-track: P1 hotfixes follow streamlined change advisory board approval, documented retrospectively.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "7.6 Rollback Strategy", level=2)
    for b in [
        "Immutable deployments: Previous version remains available for rapid reactivation.",
        "Database migration versioning: All schema changes include a corresponding rollback script.",
        "Feature flags: New functionality deployed behind feature flags, disableable without code deployment.",
        "Clear rollback criteria: Defined criteria for when rollback is appropriate vs. forward-fix.",
        "Rollback time target: Application rollback within 5 minutes for containerised Kubernetes deployments.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_heading(doc, "7.7 Data Protection During Debugging", level=2)
    for b in [
        "No raw PII in logs: Sensitive fields masked before writing to log storage.",
        "Restricted log access: Access-controlled. All access itself logged.",
        "Sanitised debug data: Production data anonymised or replaced with synthetic equivalents for debugging.",
        "Encrypted log storage: Log storage at rest encrypted. Log transmission uses TLS.",
        "Incident data handling policy: Documents how data accessed during investigation must be handled and disposed of.",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # ======================================================================
    # SECTION 8 — Governance
    # ======================================================================
    add_heading(doc, "8. Governance Model for AI-Assisted Development", level=1)

    gov_sections = [
        ("8.1 AI-Assisted Development Guidelines", [
            "Only approved AI coding tools (Claude Code) are used for production code development.",
            "Developers must provide Claude Code with architecture constraints, coding standards, security requirements, and business domain context.",
            "Architecture decisions, security design, and business logic interpretation remain exclusively human responsibilities.",
            "No production data, customer PII, or credentials included in AI prompts.",
            "All AI-generated code treated as a draft and subject to the same review, testing, and validation processes as human-authored code.",
        ]),
        ("8.2 Mandatory Human Code Review", [
            "Human review is non-negotiable for all code entering protected branches.",
            "Functional correctness, security review, architecture compliance, test coverage, and documentation are assessed.",
            "Two-person review required for: authentication/authorisation code, payment processing, encryption routines, and CI/CD pipeline configuration.",
        ]),
        ("8.3 Security Audit", [
            "Pre-release security assessment by a qualified security engineer or third-party penetration tester before each major release.",
            "Quarterly security review of access controls, secret rotation status, and dependency vulnerabilities.",
            "Continuous SAST: Security hotspots require explicit review and sign-off before merge.",
            "Annual penetration test by independent third-party security firm with findings tracked to remediation.",
        ]),
        ("8.4 Automated CI/CD Validation", [
            "Every code change must pass: linting, type safety check, unit tests (≥80% coverage), integration tests, SAST, dependency scan, and build validation.",
            "Merge to main triggers: staging deployment, E2E tests, DAST scan on staging, and performance regression test.",
            "Production deployment requires explicit approval gate, followed by smoke tests and synthetic monitoring validation.",
        ]),
        ("8.5 Version Control Governance", [
            "Trunk-based development with short-lived feature branches (merged within 2 days).",
            "main and production branches protected — no direct pushes; all changes via pull request.",
            "All commits GPG-signed; unsigned commits rejected by branch protection rules.",
            "Semantic versioning (MAJOR.MINOR.PATCH) for all production releases.",
            "Conventional Commits specification enforced to enable automated changelog generation.",
        ]),
        ("8.6 Documentation Standards", [
            "Architecture Decision Records (ADRs): Every significant architectural decision documented in docs/architecture/decisions/.",
            "API documentation: All REST APIs documented using OpenAPI 3.0, generated from code annotations.",
            "Operational runbooks: Reviewed and updated after every incident.",
            "Developer onboarding guide covering environment setup, architecture overview, workflow, and key contacts.",
            "Data dictionary: All entities and fields documented and maintained in alignment with the database schema.",
            "Changelog: CHANGELOG.md maintained per release, auto-generated from Conventional Commits.",
        ]),
    ]
    for sec_title, sec_bullets in gov_sections:
        add_heading(doc, sec_title, level=2)
        for b in sec_bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    doc.add_page_break()

    # ======================================================================
    # APPENDIX A — Glossary
    # ======================================================================
    add_heading(doc, "Appendix A: Glossary", level=1)
    add_table(doc,
        ["Term", "Definition"],
        [
            ["SAST",    "Static Application Security Testing — analysis of source code for security vulnerabilities without executing the code"],
            ["DAST",    "Dynamic Application Security Testing — testing of a running application for security vulnerabilities"],
            ["RBAC",    "Role-Based Access Control — access control model where permissions are assigned to roles, and users to roles"],
            ["OWASP",   "Open Web Application Security Project — publishes security standards and vulnerability classifications"],
            ["JWT",     "JSON Web Token — compact, URL-safe token format used for authentication and information exchange"],
            ["CI/CD",   "Continuous Integration / Continuous Deployment — automated pipeline for building, testing, and deploying software"],
            ["IaC",     "Infrastructure as Code — managing and provisioning infrastructure through machine-readable configuration files"],
            ["ADR",     "Architecture Decision Record — document capturing an important architectural decision and its context"],
            ["MTTR",    "Mean Time to Recovery — average time taken to restore service after an incident"],
            ["PII",     "Personally Identifiable Information — data that can be used to identify an individual"],
            ["mTLS",    "Mutual TLS — authentication protocol where both client and server authenticate each other using certificates"],
            ["SBOM",    "Software Bill of Materials — formal record of all software components in an application"],
        ]
    )
    doc.add_paragraph()

    # ======================================================================
    # APPENDIX B — Tooling
    # ======================================================================
    add_heading(doc, "Appendix B: Recommended Tooling Reference", level=1)
    add_table(doc,
        ["Category", "Recommended Tools"],
        [
            ["AI Coding Assistant",        "Claude Code (Anthropic)"],
            ["Static Analysis",            "SonarQube, ESLint, Semgrep, TypeScript strict mode"],
            ["Security Scanning",          "OWASP ZAP, Snyk, npm audit, Dependabot"],
            ["Testing",                    "Jest/Vitest (unit), Playwright/Cypress (E2E), k6 (load)"],
            ["Containerisation",           "Docker, Kubernetes (EKS/AKS/GKE)"],
            ["Infrastructure as Code",     "Terraform, AWS CloudFormation"],
            ["Monitoring & Observability", "Datadog, New Relic, OpenTelemetry, Prometheus/Grafana"],
            ["Log Management",             "ELK Stack, Splunk, AWS CloudWatch"],
            ["Secret Management",          "HashiCorp Vault, AWS Secrets Manager, Azure Key Vault"],
            ["API Documentation",          "Swagger / OpenAPI 3.0"],
            ["Version Control",            "Git (GitHub/GitLab/Bitbucket) with branch protection"],
        ]
    )

    doc.add_paragraph()
    add_para(doc,
        "This document is maintained by the Enterprise Architecture team. For updates or clarifications, "
        "raise a request through the standard technical documentation process.",
        italic=True)
    add_para(doc, "Document Version: 1.0  |  Review Cycle: Annually or upon significant process change", italic=True)

    # ── Save ────────────────────────────────────────────────────────────
    output_path = "/home/user/PAS_POC/docs/AI_ASSISTED_DEVELOPMENT_GUIDELINES.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    build_document()
